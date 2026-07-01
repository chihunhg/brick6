# -*- coding: utf-8 -*-
"""Generate Word doc: Gemini AI 產業別與規範擴充說明"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "Gemini-AI產業別與規範擴充說明.docx"
OUT_ALT = Path(__file__).resolve().parent / "Gemini-AI產業別與規範擴充說明-更新.docx"

FONT_NAME = "Microsoft JhengHei"
TITLE_SIZE = Pt(22)
HEADING1_SIZE = Pt(16)


def _set_run_font(run, size=None, bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size=Pt(11))
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=Pt(11))


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(
            run,
            size=HEADING1_SIZE if level == 1 else None,
            bold=True if level == 1 else None,
        )
    return h


def add_doc_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size=TITLE_SIZE, bold=True)


def build_doc():
    doc = Document()
    add_doc_title(doc, "Gemini AI 產業別與規範擴充說明")
    add_para(doc, "brick6 後台 CKEditor 產文 API（generate_editor.php）")
    add_para(doc, "文件版本：依專案現況整理｜適用 PHP 8.4 + google-gemini-php/client")

    heading(doc, "一、架構概覽", level=1)
    add_para(doc, "產業規範採單一真相來源（Single Source of Truth）設計，核心邏輯集中於 include/gemini_editor_helpers.php。")
    add_bullets(doc, [
        "前端傳入 industry 參數（POST 或 JSON body）",
        "manage/generate_editor.php 接收並驗證 prompt、source_url",
        "gemini_normalize_industry()：將中英文別名統一為 canonical key",
        "gemini_industry_rules()：依產業回傳法規與寫作規範文字",
        "gemini_editor_system_instruction()：組合 System Instruction（格式、來源限制、產業規範）",
        "呼叫 gemini-2.5-flash，回傳 JSON 欄位 html_content",
        "gemini_sanitize_editor_html()：過濾允許的 HTML 標籤，移除 Markdown 程式碼區塊",
    ])
    add_para(doc, "重點：generate_editor.php 通常不需修改；只要在 helper 新增產業定義，API 即自動生效。")

    heading(doc, "二、核心檔案對照", level=1)
    add_table(doc, ["檔案路徑", "用途"], [
        ["include/gemini_editor_helpers.php", "產業別名、規範文字、HTML 白名單、System Instruction（主要維護點）"],
        ["manage/generate_editor.php", "API 端點：接收參數、呼叫 Gemini、輸出 html_content"],
        ["include/gemini_client.php", "Gemini Client 建立與 SSL 設定（WAMP 本機）"],
        ["manage/js/editor-ai.js", "後台 detail 頁 AJAX；讀取 data-editor-industry 傳給 API"],
        ["manage/_detail_ckeditor_ai_button.php", "可重複使用的「AI 產生內容」按鈕 partial"],
        ["manage/test_ai_editor.php", "Tailwind 測試頁（產業下拉、參考網址、CKEditor 預覽）"],
    ])

    heading(doc, "三、新增一個產業（三步驟）", level=1)

    heading(doc, "步驟 1：gemini_normalize_industry() 加別名", level=2)
    add_para(doc, "在 include/gemini_editor_helpers.php 的 match 中新增 canonical key 與別名：")
    add_code(doc, """return match ($key) {
    // ...既有產業...
    'automotive', 'auto', '汽車', '車用' => 'automotive',
    default => 'general',
};""")
    add_para(doc, "建議：canonical key 一律使用英文底線（如 automotive）；中文名稱只放在別名中。")

    heading(doc, "步驟 2：gemini_industry_rules() 加規範文字", level=2)
    add_para(doc, "在同一檔案的 match 中新增對應 case，撰寫法規、語氣、格式與強制警語：")
    add_code(doc, """'automotive' => "【汽車產業規範】：\\n"
    . "- 禁止未經認證的安全或油耗數據宣稱。\\n"
    . "- 規格比較須使用 <table> 呈現。\\n"
    . "- 文末強制警語：<p><small>...</small></p>",""")
    add_para(doc, "若規範要求特定 HTML 標籤，請同步確認步驟 3 的白名單已包含。")

    heading(doc, "步驟 3：擴充允許的 HTML 標籤（若需要）", level=2)
    add_para(doc, "若新產業需使用目前未開放的標籤，請同步修改以下兩個函式（必須一致）：")
    add_bullets(doc, [
        "gemini_editor_allowed_html_tags()：供 strip_tags() 過濾用",
        "gemini_editor_allowed_html_tag_names()：寫入 System Instruction 告知模型",
    ])
    add_para(doc, "若兩者不同步，模型可能產出標籤，但後端 sanitize 會將其移除。")
    add_para(doc, "目前已允許：h1–h6, p, strong, em, ul, li, small, table, thead, tbody, tr, th, td")

    heading(doc, "四、目前已內建產業別", level=1)
    add_table(doc, ["Canonical Key", "中文別名範例", "規範重點"], [
        ["medical", "醫療", "禁止療效宣稱；語氣客觀；強制 <em> 醫療警語"],
        ["biotech", "生技、保健食品", "禁止降三高／抗癌等字眼；ul/li 列專利成分"],
        ["electronics", "電子、半導體、高科技", "產業用語；規格須用 <table>"],
        ["listed_company", "上市櫃、IR", "IR 官方語氣；禁止預測股價；<small> 免責聲明"],
        ["japanese_client", "日系、日商", "職人精神、御中語氣；標題層級嚴謹"],
        ["finance", "金融", "禁止保證獲利；<small> 風險警語"],
        ["beauty", "美妝、化妝品", "禁止誇大美妝詞彙"],
        ["general", "（預設）", "一般商業規範；未知別名 fallback 至此"],
    ])

    heading(doc, "五、前端需同步的地方（選用）", level=1)
    add_table(doc, ["位置", "用途"], [
        ["manage/test_ai_editor.php", "測試頁產業下拉選單 <option>"],
        ["各模組 manage/*/_detail.php", "設定 $editorAiIndustry = '產業key'"],
        ["manage/_detail_ckeditor_ai_button.php", "按鈕 data-editor-industry 屬性（由 partial 輸出）"],
    ])
    add_para(doc, "後台 detail 頁範例：")
    add_code(doc, """$editorAiFieldId = 'Contents1_' . $i;
$editorAiIndustry = 'automotive';
require dirname(__DIR__) . '/_detail_ckeditor_ai_button.php';""")
    add_para(doc, "editor-ai.js 不需修改，它會將 data-editor-industry 原樣 POST 給 API。")

    heading(doc, "六、修改既有產業規範", level=1)
    add_bullets(doc, [
        "僅修改 gemini_industry_rules() 中對應 case 的文字即可",
        "不必修改 generate_editor.php 或 Gemini 連線程式",
        "若變更強制警語的 HTML 結構，請確認 allowed HTML 白名單包含相關標籤",
    ])

    heading(doc, "七、API 請求參數摘要", level=1)
    add_table(doc, ["參數", "必填", "說明"], [
        ["prompt", "是", "寫作任務提示詞"],
        ["source_url", "是", "參考網址（http/https）；模型僅能依此改寫，禁止捏造"],
        ["industry", "否", "產業 key；預設 general"],
    ])
    add_para(doc, "成功回應範例：")
    add_code(doc, '{"html_content":"<h2>...</h2><p>...</p>"}')

    heading(doc, "八、維護建議", level=1)
    add_bullets(doc, [
        "一個產業 = 一個 match case：規範、語氣、格式、警語寫在同一區塊，方便法遵 review",
        "先用 manage/test_ai_editor.php 驗證：選產業、填真實 source_url、檢查 CKEditor 與 JSON",
        "別名集中於 gemini_normalize_industry()，避免前端各處使用不同字串",
        "未知產業會 fallback 到 general，不會報錯；若需嚴格管控可於 API 加白名單驗證",
        "環境變數：.env 設定 GEMINI_API_KEY；本機 WAMP 可設 GEMINI_SSL_VERIFY=0",
    ])

    heading(doc, "九、未來進階重構（選用）", level=1)
    add_para(doc, "目前規範寫在 PHP match 中，適合約 10 個以內、偶爾調整的場景。若產業數量增多或需非法遵人員維護，可考慮：")
    add_bullets(doc, [
        "抽至 config/industry_rules.php 或 JSON/YAML 設定檔",
        "或建立資料庫 industry_rules 表 + 後台 CRUD 管理介面",
    ])
    add_para(doc, "現階段建議維持 helper 集中管理，待需求明確再重構。")

    heading(doc, "十、總結", level=1)
    add_para(
        doc,
        "擴充產業 = 修改 include/gemini_editor_helpers.php 的 gemini_normalize_industry() "
        "與 gemini_industry_rules()；必要時更新 HTML 白名單；測試頁與 detail 頁補上選項。"
        "generate_editor.php 與 Gemini 連線層通常無需變動。",
        bold=True,
    )

    return doc


def main():
    doc = build_doc()
    try:
        doc.save(OUT)
        print("Saved:", OUT)
    except PermissionError:
        doc.save(OUT_ALT)
        print("Saved (alt, close original docx first):", OUT_ALT)


if __name__ == "__main__":
    main()
