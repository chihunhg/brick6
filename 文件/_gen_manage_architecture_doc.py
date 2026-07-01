# -*- coding: utf-8 -*-
"""Generate Word doc: manage/ 後台程式架構說明與 2 小時教學大綱"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parent / "manage後台程式架構說明-教學大綱.docx"
OUT_ALT = Path(__file__).resolve().parent / "manage後台程式架構說明-教學大綱-更新.docx"

FONT_NAME = "Microsoft JhengHei"
HEADING1_SIZE = Pt(16)
TITLE_SIZE = Pt(22)


def _set_run_font(run, size=None, bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size=Pt(11))
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=Pt(11))


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, size=HEADING1_SIZE if level == 1 else None, bold=True if level == 1 else None)
    return h


def add_doc_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size=TITLE_SIZE, bold=True)
    return p


def add_doc_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    _set_run_font(run, size=HEADING1_SIZE, bold=True)
    return p


def section_overview(doc):
    heading(doc, "文件說明", level=1)
    add_para(
        doc,
        "本文件供程式開發人員了解 brick6 專案 manage/ 後台管理系統之程式架構、"
        "功能目錄分工、共用程式用途，以及預計 2 小時之內部教學大綱。"
        "後台採 PHP 8.4 + PDO + 模組化 CRUD 框架，以資料庫 module_p / module_d 驅動選單與權限，"
        "各內容單元以獨立目錄搭配共用 _*.php 元件組裝而成。",
    )
    add_table(
        doc,
        ["核心概念", "說明"],
        [
            ["manNo", "對應 module_p.PKey，主單元編號（網址參數與 Session）"],
            ["subNo", "對應 module_d.PKey，子單元（多層分類時）"],
            ["strLink", "module_p 中的目錄名，如 news → manage/news/"],
            ["FunctionID", "Session 內可存取模組 ID 清單；非 Admin 須通過 _module.php 檢查"],
        ],
    )


def section_architecture(doc):
    heading(doc, "一、整體架構", level=1)

    heading(doc, "1.1 進入流程", level=2)
    add_bullets(
        doc,
        [
            "manage/index.php：檢查 Session，導向 login/index.php 或 login/login.php",
            "manage/login/index.php：登入表單、驗證帳密、寫入 Session（Manage、Login_ID、FunctionID）",
            "manage/login/login.php：登入後首頁（Hello 畫面與操作提示）",
        ],
    )

    heading(doc, "1.2 Bootstrap 載入鏈", level=2)
    add_para(doc, "每個後台頁面（除純登入頁外）通常以 require '../_inc.php' 開始：", bold=True)
    add_bullets(
        doc,
        [
            "include/host.php → 載入 .env、網域與路徑常數",
            "include/Conn.php → 全域 $pdo 資料庫連線",
            "include/dbclass.php → recordset（查詢）、dbPDO（CRUD）",
            "include/Function.php、log.php、sec.php → 工具、操作紀錄、安全標頭",
            "include/common.php → 級聯刪除、載入 crud_helpers.php",
            "include/image.php → 縮圖與 WebP",
            "manage/_lang_slots.php、_form_bag.php、_child_helpers.php 等後台專用 helper",
        ],
    )

    heading(doc, "1.3 模組權限（_module.php）", level=2)
    add_bullets(
        doc,
        [
            "依 manNo 讀取 module_p，取得 Module_Name、intLayer、intList 等",
            "依 subNo 讀取 module_d，設定子單元標題",
            "比對 $_SESSION['FunctionID'] 與 $Module_PKey；Admin 帳號例外",
            "設定 $GLOBALS 供列表、表單、麵包屑使用",
        ],
    )

    heading(doc, "1.4 標準請求流程", level=2)
    add_table(
        doc,
        ["頁型", "流程摘要"],
        [
            ["列表 list.php", "_inc → _module → _config → crud_process_list_actions → 查詢分頁 → layout → _list.php"],
            ["新增/編輯 add.php / update.php", "_inc → _module → _config → _form_data → _detail.php → POST 至 addin.php"],
            ["寫入 addin.php", "CSRF 驗證 → 欄位驗證 → 上傳 → crud_upsert_master → 子表 → _return_list.php"],
        ],
    )


def section_shared_manage(doc):
    heading(doc, "二、manage/ 根目錄共用檔", level=1)
    add_para(doc, "開發新單元時主要「引用」以下檔案，而非複製邏輯：")

    add_table(
        doc,
        ["檔案", "用途"],
        [
            ["_inc.php", "後台 Bootstrap：Session、環境、DB、語系、登入檢查、CSRF、上傳目錄"],
            ["_module.php", "模組標題、Layer、FunctionID 權限檢查"],
            ["_layout_head.php / _layout_body_open.php / _layout_body_close.php", "HTML 外框"],
            ["_header.php / _sidebar.php / _footer.php", "頂欄、側欄選單、頁尾"],
            ["_breadcrumbs.php", "麵包屑導覽"],
            ["_in_code_head.php / _in_javascript.php / _in_code_bottom.php", "CSS/JS 引入與頁尾腳本"],
            ["_select.php", "列表工具列（全選、批次刪除/上下架、新增）"],
            ["_search.php / _search_list.php", "關鍵字與篩選 UI"],
            ["_page.php", "分頁元件"],
            ["_submit.php", "表單底部 hidden 欄位與送出/關閉按鈕"],
            ["_return_list.php", "儲存成功後帶搜尋條件導回列表"],
            ["_upload_endpoint.php", "列表單筆上下架 AJAX 共用端點"],
            ["_del_img_endpoint.php", "刪除圖片共用端點"],
            ["_child_helpers.php / _child_list.php / _child_form.php", "子模組（如 album_d）列表/表單共用"],
            ["_form_bag.php", "表單變數 bag 模式（$GLOBALS['xxx_form_vars']）"],
            ["_lang_slots.php / _detail_lang_seo_fields.php", "多語系欄位與 SEO 欄位片段"],
            ["_ckeditor.php", "富文字編輯器嵌入"],
            ["_tag_relation_block.php", "標籤關聯 UI"],
            ["_detail_config.sample.php", "_config.php 範本與 addin 開發指引"],
        ],
    )


def section_modules(doc):
    heading(doc, "三、功能模組目錄", level=1)

    heading(doc, "3.1 標準模組檔案命名", level=2)
    add_table(
        doc,
        ["檔案", "角色"],
        [
            ["_config.php", "資料表對應（master / lang / msg / img / fk / csrf）"],
            ["list.php", "列表入口：查詢、分頁、批次操作、layout"],
            ["_list.php", "列表 HTML 列（模組專屬欄位）"],
            ["add.php / update.php", "新增/編輯表單頁"],
            ["addin.php", "POST 處理：驗證 → 上傳 → 寫 DB → 導回"],
            ["_detail.php", "表單 HTML"],
            ["_form_data.php", "從 DB 載入/初始化表單資料"],
            ["_upload.php", "定義 MANAGE_UPLOAD_MODULE_DIR 後引入 _upload_endpoint.php"],
            ["_del_img.php", "刪圖端點包裝"],
        ],
    )

    heading(doc, "3.2 現有模組目錄一覽", level=2)
    add_table(
        doc,
        ["目錄", "業務用途", "備註"],
        [
            ["login/", "登入、登出、後台首頁", "不經 _module.php 權限樹"],
            ["control/", "權限帳號、SEO 設定(webset)、改密碼", "系統管理"],
            ["module/", "單元設定（module_p 本身的管理）", "僅 Admin"],
            ["class1/ class2/ class3/", "三層分類", "多數內容模組會引用"],
            ["news/", "最新消息", "標準 CRUD 範本"],
            ["paper/", "刊物/文件", "同 news 結構"],
            ["product/", "產品", "含關聯、autocomplete"],
            ["video/", "影音", ""],
            ["faq/ knowledge/", "FAQ、知識庫", ""],
            ["investor/", "投資人訊息", ""],
            ["ad/", "廣告/Banner", ""],
            ["album/", "相簿主檔", "父模組"],
            ["album_d/", "相簿照片明細", "子模組（依父 Album_PKey）"],
            ["weblink/", "連結管理", ""],
            ["filedown/", "檔案下載", ""],
            ["company/", "公司/據點", ""],
            ["member/", "會員", ""],
            ["tag/", "標籤", ""],
            ["question/ question_class/ question_item/", "問卷系統", "多表關聯"],
        ],
    )

    heading(doc, "3.3 支援目錄", level=2)
    add_table(
        doc,
        ["目錄", "用途"],
        [
            ["ajax/", "JSON API（如分類連動 ajax.php）"],
            ["elFinder/", "CKEditor 檔案總管"],
            ["css/ js/", "後台樣式與互動（script.js、tag-relation.js、file_preview.js）"],
            ["ckeditor/", "編輯器設定"],
        ],
    )


def section_include(doc):
    heading(doc, "四、include/ 共用函式庫", level=1)
    add_table(
        doc,
        ["檔案", "用途"],
        [
            ["host.php", "載入 .env、路徑常數"],
            ["Conn.php", "全域 $pdo 連線"],
            ["dbclass.php", "recordset（查詢）、dbPDO（CRUD）"],
            ["common.php", "級聯刪除 delete_cascade_by_ids、載入 crud_helpers"],
            ["crud_helpers.php", "CRUD 核心：列表、分頁、驗證、上傳、多語系儲存、模組設定"],
            ["sec.php", "e() XSS 轉義、CSRF、CSP、安全 redirect"],
            ["Function.php", "SqlFilter、location_href 等通用工具"],
            ["log.php", "manage_history、sql_error"],
            ["image.php", "縮圖、WebP"],
            ["tag_relation_helpers.php", "標籤關聯存取"],
            ["json_response.php", "AJAX 統一 JSON 回應"],
            ["frontend_modules.php", "前台 slug 與 module_p.PKey 對照"],
            ["frontend_helpers.php", "前台列表、明細、分類共用函式"],
        ],
    )


def section_crud_helpers(doc):
    heading(doc, "五、常用 CRUD Helper 速查", level=1)
    add_table(
        doc,
        ["函式", "用途"],
        [
            ["manage_detail_set_config() / manage_detail_tables()", "註冊並讀取模組表設定"],
            ["crud_cfg() / crud_process_list_actions()", "列表刪除設定與批次處理"],
            ["crud_module_where()", "依 Module_PKey 組 WHERE"],
            ["crud_paginate() / crud_fetch_all()", "分頁查詢"],
            ["crud_upsert_master()", "新增/更新主檔"],
            ["crud_save_lang_slots()", "多語系子表"],
            ["crud_save_msg_blocks_multilang()", "多語系 HTML 內文"],
            ["crud_save_img_slots()", "多圖槽位"],
            ["crud_form_error_redirect()", "驗證失敗導回並顯示錯誤"],
            ["crud_addin_return_url()", "組 addin 失敗/成功返回 URL"],
            ["delete_cascade_by_ids()", "刪主檔 + 子表 + 實體檔案"],
            ["manage_module_show_detail_field()", "依 Module_PKey 開關首頁/簡述/列表圖/標籤等共用欄位"],
        ],
    )

    heading(doc, "5.1 _config.php 範例（news）", level=2)
    add_para(
        doc,
        "return [\n"
        "    'master' => 'news',\n"
        "    'img' => 'news_img', 'lang' => 'news_lang', 'msg' => 'news_msg',\n"
        "    'fk' => 'News_PKey', 'module_pk_col' => 'Module_PKey',\n"
        "    'csrf' => 'news_addin', 'tag_relation_parent_col' => 'News_PKey',\n"
        "];",
    )


def section_security(doc):
    heading(doc, "六、權限與安全機制", level=1)
    add_bullets(
        doc,
        [
            "Session 登入：_inc.php 檢查 $_SESSION['Manage'] === 'Yes'",
            "模組權限：_module.php 比對 FunctionID 與 Module_PKey（Admin 例外）",
            "CSRF：列表 crud_csrf_guard_list()；表單 crud_csrf_ensure_page() / crud_csrf_verify_form()",
            "SQL 注入：recordset / dbPDO 具名參數；crud_is_safe_sql_identifier() 防表名注入",
            "XSS：輸出使用 e()（htmlspecialchars）",
            "操作紀錄：manage_history() 寫入 managelog",
            "側欄選單：_sidebar.php 讀 module_p，分為網站管理/首頁管理/單元管理/系統管理",
        ],
    )


def section_child_module(doc):
    heading(doc, "七、子模組模式（Parent / Child）", level=1)
    add_para(
        doc,
        "相簿為典型範例：album/ 為相簿主檔，album_d/ 為某相簿下的照片列表，"
        "透過 manage_child_list_prepare() 解析父鍵 Album_PKey。"
        "子模組使用 manage_child_return_url() 等 helper，列表/表單邏輯與主模組分離但共用 CRUD helper。",
    )
    add_table(
        doc,
        ["一般內容單元", "子模組（manage_child_*）"],
        [
            ["add.php 自寫初始化", "manage_child_form_add_prepare()"],
            ["update.php 自寫 load", "manage_child_form_update_prepare()"],
            ["addin.php 自寫", "manage_child_addin_run() 或 _form_data 內 validate/save"],
            ["list.php", "manage_child_list_prepare() + manage_child_list_render()"],
        ],
    )


def section_checklist(doc):
    heading(doc, "八、新增模組快速 Checklist", level=1)
    add_bullets(
        doc,
        [
            "在 DB 建立主檔與子表（lang/msg/img 視需求）",
            "在 module_p 新增一筆，strLink = 目錄名",
            "建立 manage/{name}/，複製 news/ 或參考 _detail_config.sample.php",
            "修改 _config.php 表名與 FK",
            "調整 _list.php、_detail.php 欄位 UI",
            "調整 _form_data.php 載入邏輯",
            "確認 addin.php 驗證與 $data_array 欄位",
            "將模組 PKey 加入測試帳號的 FunctionID",
            "在 _inc.php 的 $manage_module_detail_fields 設定共用欄位開關",
            "前台頁面於 include/frontend_modules.php 登錄 slug 對應 PKey",
        ],
    )


def section_curriculum(doc):
    heading(doc, "九、2 小時教學大綱", level=1)
    add_para(
        doc,
        "以下為程式開發人員內訓建議時程，合計約 120 分鐘。"
        "每段結束可安排 5 分鐘 Q&A（未計入下列時間）。",
    )

    heading(doc, "第 1 部分：環境與架構鳥瞰（25 分鐘）", level=2)
    add_table(
        doc,
        ["時間", "主題", "內容"],
        [
            ["0:00–0:10", "專案定位", "前後台關係、manage/ vs 根目錄前台 PHP、module_p.strLink 對應"],
            ["0:10–0:20", "進入流程", "index.php → 登入 → login/login.php；Session 欄位說明"],
            ["0:20–0:25", "目錄地圖", "共用 _*.php、功能模組目錄、include/ 分工"],
        ],
    )
    add_para(doc, "演示：瀏覽器走一遍「登入 → 側欄 → news 列表 → 新增」。")

    heading(doc, "第 2 部分：Bootstrap 與資料存取（25 分鐘）", level=2)
    add_table(
        doc,
        ["時間", "主題", "內容"],
        [
            ["0:25–0:35", "_inc.php 深度", ".env/PDO、語系、webset、$upload_folder、登入閒置登出"],
            ["0:35–0:45", "_module.php", "manNo/subNo、Layer、FunctionID 權限拒絕流程"],
            ["0:45–0:50", "資料層", "recordset vs dbPDO、錯誤寫 log 慣例"],
        ],
    )
    add_para(doc, "練習：追蹤 list.php?manNo=X&subNo=Y 如何得到 $Module_Name。")

    heading(doc, "第 3 部分：標準模組 CRUD 拆解（35 分鐘）", level=2)
    add_table(
        doc,
        ["時間", "主題", "內容"],
        [
            ["0:50–1:00", "_config.php", "master/lang/msg/img/fk/csrf 各欄位意義；複製 sample 建模組"],
            ["1:00–1:10", "列表頁", "list.php + _list.php + _select.php + _search.php + 分頁"],
            ["1:10–1:20", "表單頁", "add.php/update.php → _form_data.php → _detail.php → _submit.php"],
            ["1:20–1:25", "寫入", "addin.php 驗證→上傳→upsert→子表→_return_list.php"],
        ],
    )
    add_para(doc, "演示：對照 news/ 四支主檔，畫出 POST 資料流。")

    heading(doc, "第 4 部分：進階主題（25 分鐘）", level=2)
    add_table(
        doc,
        ["時間", "主題", "內容"],
        [
            ["1:25–1:35", "多語系與 SEO", "_lang_slots.php、*_lang 表、crud_save_lang_slots、_detail_lang_seo_fields.php"],
            ["1:35–1:42", "圖片上傳", "crud_upload_file_slots、_upload_endpoint.php、_del_img_endpoint.php"],
            ["1:42–1:50", "子模組", "album_d + manage_child_list_prepare 父鍵解析"],
        ],
    )

    heading(doc, "第 5 部分：安全、AJAX 與擴充實務（10 分鐘）", level=2)
    add_table(
        doc,
        ["時間", "主題", "內容"],
        [
            ["1:50–1:55", "安全 checklist", "CSRF、SqlFilter、e()、CSP、manage_history"],
            ["1:55–2:00", "擴充指南", "新增模組步驟：module_p 註冊 → 複製 news → 改 _config → 前台接資料"],
        ],
    )

    heading(doc, "課後作業（選做）", level=2)
    add_bullets(
        doc,
        [
            "複製 news/ 建立練習模組，只保留標題 + 上下架欄位",
            "閱讀 ajax/ajax.php，理解 Class2/Class3 連動 JSON 格式",
            "追蹤 delete_cascade_by_ids() 刪除一筆 news 時刪了哪些表與檔案",
        ],
    )


def section_references(doc):
    heading(doc, "十、延伸閱讀", level=1)
    add_bullets(
        doc,
        [
            "manage/_detail_config.sample.php — _config.php 範本與 addin 註解",
            "文件/後台新增與調整單元-修改程式清單與注意事項.docx — 新增/調整單元詳細清單",
            "include/frontend_modules.php — 前台 slug 與 PKey 對照",
            "側欄「後台操作教學」YouTube 播放清單 — 操作面教學影片",
        ],
    )


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_doc_title(doc, "manage/ 後台程式架構說明")
    add_doc_subtitle(doc, "功能目錄、共用程式與 2 小時教學大綱")

    section_overview(doc)
    section_architecture(doc)
    section_shared_manage(doc)
    section_modules(doc)
    section_include(doc)
    section_crud_helpers(doc)
    section_security(doc)
    section_child_module(doc)
    section_checklist(doc)
    section_curriculum(doc)
    section_references(doc)

    try:
        doc.save(str(OUT))
        print("Saved:", OUT)
    except PermissionError:
        doc.save(str(OUT_ALT))
        print("Saved (alt, close original docx first):", OUT_ALT)


if __name__ == "__main__":
    main()
