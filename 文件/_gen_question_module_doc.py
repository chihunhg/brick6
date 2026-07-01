# -*- coding: utf-8 -*-
"""Generate Word doc: Question 問卷模組檔案對照表"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parent / "Question問卷模組-檔案對照表.docx"
OUT_ALT = Path(__file__).resolve().parent / "Question問卷模組-檔案對照表-更新.docx"

FONT_NAME = "Microsoft JhengHei"
HEADING1_SIZE = Pt(16)
TITLE_SIZE = Pt(22)


def _set_run_font(run, size=None, bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    _set_run_font(run, size=Pt(9))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size=Pt(11))
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=Pt(11))


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, size=HEADING1_SIZE if level == 1 else None, bold=True if level == 1 else None)


def add_doc_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size=TITLE_SIZE, bold=True)


def add_doc_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    _set_run_font(run, size=HEADING1_SIZE, bold=True)


def section_overview(doc):
    heading(doc, "一、整體架構與導覽流程", level=1)
    add_para(
        doc,
        "問卷後台採三層結構：主檔（question）→ 類別（question_class）→ 題目（question_item，含答案選項）。"
        "另有填答匯出（output.php）與前台填寫頁（questionnaire.php）。",
    )
    add_table(
        doc,
        ["層級", "目錄", "業務意義", "父鍵參數"],
        [
            ["L1", "manage/question/", "問卷主題（代號、期間、收件信箱、介紹文）", "manNo / subNo"],
            ["L2", "manage/question_class/", "問卷內的分類區塊", "Question_PKey"],
            ["L3", "manage/question_item/", "各類別下的題目與選項", "Question_D_PKey（類別 PKey）"],
        ],
    )
    add_para(doc, "後台操作路徑：", bold=True)
    add_bullets(
        doc,
        [
            "問卷列表 → [設定] → 類別列表 → [設定] → 題目列表",
            "問卷列表 → [編輯] → 問卷表單（add.php / update.php）",
            "問卷列表 → [匯出(n)] → output.php（新視窗 Excel）",
        ],
    )


def section_tables(doc):
    heading(doc, "二、資料表對照", level=1)

    heading(doc, "2.1 問卷主檔（manage/question/）", level=2)
    add_table(
        doc,
        ["資料表", "用途", "外鍵"],
        [
            ["question", "主檔：strNo、strName、EMail、OpenDate、EndDate、Sort、Upload、Module_PKey", "—"],
            ["question_lang", "多語系標題、SEO 等", "Question_PKey"],
            ["question_msg", "各語系問卷介紹 HTML（Sort = 語系序）", "Question_PKey"],
            ["question_img", "列表圖（1 槽，依 list 欄位開關）", "Question_PKey"],
        ],
    )

    heading(doc, "2.2 類別（manage/question_class/）", level=2)
    add_table(
        doc,
        ["資料表", "用途", "外鍵"],
        [
            ["question_class", "類別主檔：Sort、strName、Upload", "Question_PKey"],
            ["question_class_lang", "類別多語系名稱", "Question_Class_PKey"],
        ],
    )

    heading(doc, "2.3 題目與答案（manage/question_item/）", level=2)
    add_table(
        doc,
        ["資料表", "用途", "外鍵"],
        [
            ["question_item", "題目：Qtype、Other、Must、Sort、Upload", "Question_PKey、Question_D_PKey"],
            ["question_itme_lang", "題目多語系（表名為歷史拼字 itme）", "Question_Item_PKey"],
            ["question_answer", "選項（單選/複選，最多 10 槽）", "Question_I_PKey"],
            ["question_answer_lang", "選項多語系", "Question_Answer_PKey"],
        ],
    )

    heading(doc, "2.4 填答紀錄", level=2)
    add_table(
        doc,
        ["資料表", "用途"],
        [
            ["question_report_p", "填答主檔（一筆問卷送出）"],
            ["question_report_d", "填答明細（各題答案）"],
            ["question_report", "舊版／相容填答表（刪除時一併處理）"],
        ],
    )

    heading(doc, "2.5 檢視用 View", level=2)
    add_table(
        doc,
        ["View", "用途"],
        [
            ["view_question", "前台列表／問卷頁讀取"],
            ["view_question_class", "類別（含語系）"],
            ["view_question_item", "匯出 Excel 題目欄位"],
        ],
    )


def section_question_files(doc):
    heading(doc, "三、manage/question/ 檔案對照（L1 主檔）", level=1)
    add_table(
        doc,
        ["檔案", "類型", "職責", "備註"],
        [
            ["_config.php", "設定", "master/lang/msg/img、csrf、has_sort、content_blocks=1", "list_csrf=question_list"],
            ["_helpers.php", "共用", "顯示名稱、級聯刪除、msg 儲存、填答筆數", "被 class/item 引用"],
            ["_form_data.php", "資料", "class1_detail_* 包裝（沿用 class1 命名）", "載入/預設/複製"],
            ["list.php", "入口", "列表、搜尋、刪除前 question_delete_related_rows", ""],
            ["_list.php", "視圖", "代號、主題、上下架、問卷分類、匯出", "連 question_class/list"],
            ["add.php", "入口", "新增；crud_next_sort；可複製", "→ _detail.php"],
            ["update.php", "入口", "編輯", "→ _detail.php"],
            ["_detail.php", "視圖", "代號、信箱、期間、多語標題、介紹 CKEditor", "list 欄位開關"],
            ["addin.php", "寫入", "驗證 EMail、介紹 b64；question_save_msg_langs", "非標準 crud_save_msg_blocks"],
            ["_upload.php", "API", "列表上下架", "→ _upload_endpoint.php"],
            ["_del_img.php", "API", "刪列表圖", "→ _del_img_endpoint.php"],
            ["output.php", "匯出", "PhpSpreadsheet 匯出填答 Excel", "僅 require _inc.php"],
        ],
    )
    add_para(doc, "addin.php 與一般 news 的差異：", bold=True)
    add_bullets(
        doc,
        [
            "必填 EMail（收件信箱）",
            "問卷介紹寫入 question_msg（question_save_msg_langs；PKey = Question_PKey×100 + Sort）",
            "無 question_link 子表實際寫入（_config 的 link 僅語意對應 item 模組）",
        ],
    )


def section_class_files(doc):
    heading(doc, "四、manage/question_class/ 檔案對照（L2 子模組）", level=1)
    add_para(doc, "採 manage_child_* 子模組骨架（同 album_d 模式）。class/item 無 _upload、_del_img。")
    add_table(
        doc,
        ["檔案", "類型", "職責"],
        [
            ["_config.php", "設定", "master=question_class；lang=question_class_lang；parent_fk=Question_PKey"],
            ["_form_data.php", "資料", "父問卷解析、儲存、刪除、form bag"],
            ["list.php", "入口", "manage_child_list_prepare + manage_child_list_render"],
            ["_list.php", "視圖", "問卷主題、類別名、問卷題目按鈕 → question_item/list"],
            ["add.php / update.php", "入口", "manage_child_form_*_prepare"],
            ["_detail.php", "視圖", "類別名稱（多語）、順序、上下架"],
            ["addin.php", "寫入", "manage_child_addin_run → question_class_save_multilang"],
        ],
    )
    heading(doc, "4.1 _form_data.php 主要函式", level=2)
    add_table(
        doc,
        ["函式", "用途"],
        [
            ["question_class_resolve_question_pkey()", "從 Question_PKey / PKey 解析父問卷"],
            ["question_class_load_parent()", "載入父問卷，回傳 Question_Name"],
            ["question_class_save_multilang()", "INSERT/UPDATE 主檔 + question_class_lang"],
            ["question_class_delete_related_rows()", "刪類別 → 連動刪該類別下所有 question_item"],
            ["question_class_form_load() / form_init()", "表單 bag：$GLOBALS['question_class_form']"],
            ["question_class_next_sort()", "新類別預設順序"],
            ["question_class_display_strname()", "列表顯示類別名"],
        ],
    )
    heading(doc, "4.2 list.php prepare 選項", level=2)
    add_table(
        doc,
        ["選項", "值"],
        [
            ["parent_resolve", "question_class_resolve_question_pkey"],
            ["parent_fail_url", "../question/list.php"],
            ["list_where", "WHERE Question_PKey = :Question_PKey"],
            ["delete_handler", "question_class_delete_related_rows"],
            ["add_url", "add.php?Question_PKey={父PKey}"],
        ],
    )


def section_item_files(doc):
    heading(doc, "五、manage/question_item/ 檔案對照（L3 子模組）", level=1)
    add_table(
        doc,
        ["檔案", "類型", "職責"],
        [
            ["_config.php", "設定", "master=question_item；lang=question_itme_lang；answer_slots=10"],
            ["_form_data.php", "資料", "父類別解析、題目/答案 CRUD（最複雜）"],
            ["list.php", "入口", "manage_child_list_prepare；list_return_fk=Question_D_PKey"],
            ["_list.php", "視圖", "題型 question_type()、題目名、編輯/複製"],
            ["add.php / update.php", "入口", "manage_child_form_*_prepare"],
            ["_detail.php", "視圖", "題型、必填、答案槽 A_Name{槽}_{語系}"],
            ["addin.php", "寫入", "manage_child_addin_run → question_item_save"],
        ],
    )
    heading(doc, "5.1 題型（Qtype）", level=2)
    add_table(
        doc,
        ["值", "名稱", "需 question_answer"],
        [
            ["1", "單選題", "是"],
            ["2", "複選題", "是"],
            ["3", "單行文字題", "否"],
            ["4", "多行文字題", "否"],
            ["5", "日期題", "否"],
            ["6", "EMail", "否"],
        ],
    )
    heading(doc, "5.2 _form_data.php 主要函式", level=2)
    add_table(
        doc,
        ["函式", "用途"],
        [
            ["question_item_resolve_class_pkey()", "父鍵：Question_D_PKey"],
            ["question_item_load_parent()", "JOIN question_class + question"],
            ["question_item_save()", "寫 question_item + lang + 答案"],
            ["question_item_save_answers()", "寫 question_answer / question_answer_lang"],
            ["question_item_delete_related_rows()", "刪題目 + 答案 + lang"],
            ["question_item_row_belongs_to_class()", "編輯權限；相容 Question_D_PKey=class.Sort"],
            ["question_item_answer_filter_key($slot,$lang)", "例：A_Name3_2 = 第3選項、語系2"],
        ],
    )


def section_shared(doc):
    heading(doc, "六、共用程式與級聯刪除", level=1)
    add_table(
        doc,
        ["檔案", "被誰使用", "用途"],
        [
            ["manage/question/_helpers.php", "question / class / item", "刪除、顯示名稱、msg、匯出計數"],
            ["manage/_child_helpers.php", "class / item", "manage_child_return_url、save_lang_rows"],
            ["manage/_child_list.php", "class/item list.php", "manage_child_list_prepare/render"],
            ["manage/_child_form.php", "class/item add/update", "manage_child_form_*_prepare"],
            ["question_child_return_url()", "class/item", "包裝 manage_child_return_url"],
        ],
    )
    add_para(doc, "級聯刪除鏈（刪問卷時）：", bold=True)
    add_bullets(
        doc,
        [
            "question/list 刪除 → question_delete_related_rows",
            "→ 各 question_class → question_class_delete_related_rows",
            "→ 各 question_item → question_item_delete_related_rows",
            "→ question_delete_report_rows（填答紀錄）",
        ],
    )


def section_frontend(doc):
    heading(doc, "七、前台與後台對應", level=1)
    add_table(
        doc,
        ["檔案", "職責"],
        [
            ["questionnaire.php", "問卷填寫頁；frontend_module_pkey('question') → PKey 19"],
            ["include/frontend_modules.php", "'question' => 19"],
            [".htaccess", "questionnaire.htm → questionnaire.php"],
        ],
    )
    add_para(doc, "questionnaire.php 流程：讀 view_question → questionnaire_load_sections → POST 寫 report_p/d", bold=True)
    add_table(
        doc,
        ["後台維護", "前台使用"],
        [
            ["question Upload、OpenDate/EndDate", "是否顯示問卷"],
            ["question_msg 介紹", "頁首說明區"],
            ["question_class Sort", "區塊順序"],
            ["question_item Qtype/Must/Other", "表單欄位型態與驗證"],
            ["question_answer", "單選/複選選項"],
        ],
    )


def section_csrf_url(doc):
    heading(doc, "八、CSRF 與 URL 參數", level=1)
    add_table(
        doc,
        ["模組", "表單 csrf", "列表 csrf"],
        [
            ["question", "question_addin", "question_list"],
            ["question_class", "question_class_addin", "question_class_list"],
            ["question_item", "question_item_addin", "question_item_list"],
        ],
    )
    heading(doc, "8.1 URL 參數速查", level=2)
    add_table(
        doc,
        ["頁面", "必要參數", "說明"],
        [
            ["question/list.php", "manNo, subNo", "模組權限"],
            ["question_class/list.php", "Question_PKey 或 PKey", "哪一份問卷"],
            ["question_item/list.php", "Question_D_PKey 或 PKey", "哪一個類別"],
            ["question_class/add.php", "Question_PKey", "新增類別所屬問卷"],
            ["question_item/add.php", "Question_D_PKey", "新增題目所屬類別"],
            ["question/output.php", "PKey", "問卷主鍵（匯出）"],
            ["questionnaire.php", "PKey（可選）", "指定問卷；無則第一筆上架"],
        ],
    )


def section_notes(doc):
    heading(doc, "九、開發注意事項", level=1)
    add_bullets(
        doc,
        [
            "表名拼字：question_itme_lang（itme 非 item），改表需全專案搜尋。",
            "子模組 class/item 的 addin 用 manage_child_addin_run，不是複製 news 的 addin。",
            "問卷主檔 _form_data.php 仍用 class1_detail_* 函式名（從 class1 複製）。",
            "舊資料相容：question_item_row_belongs_to_class 允許 Question_D_PKey 存類別 Sort。",
            "output.php 較舊（直接 $_REQUEST），修改問卷 CRUD 通常不動它。",
            "列表圖受 $manage_module_detail_fields 的 list 開關控制（_inc.php）。",
        ],
    )
    heading(doc, "十、檔案總覽", level=1)
    add_table(
        doc,
        ["目錄", "檔案數", "說明"],
        [
            ["manage/question/", "12", "主檔 + 匯出 + helpers"],
            ["manage/question_class/", "8", "子模組（無 upload/del_img）"],
            ["manage/question_item/", "8", "子模組（無 upload/del_img）"],
            ["questionnaire.php", "1", "前台填寫"],
        ],
    )


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_doc_title(doc, "Question 問卷模組")
    add_doc_subtitle(doc, "檔案對照表（manage/question · question_class · question_item）")

    section_overview(doc)
    section_tables(doc)
    section_question_files(doc)
    section_class_files(doc)
    section_item_files(doc)
    section_shared(doc)
    section_frontend(doc)
    section_csrf_url(doc)
    section_notes(doc)

    try:
        doc.save(str(OUT))
        print("Saved:", OUT)
    except PermissionError:
        doc.save(str(OUT_ALT))
        print("Saved (alt):", OUT_ALT)


if __name__ == "__main__":
    main()
