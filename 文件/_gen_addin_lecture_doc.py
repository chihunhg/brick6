# -*- coding: utf-8 -*-
"""Generate Word doc: addin.php 逐步註解講義（教學第 3 部分）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parent / "addin.php逐步註解講義.docx"
OUT_ALT = Path(__file__).resolve().parent / "addin.php逐步註解講義-更新.docx"

FONT_NAME = "Microsoft JhengHei"
MONO_FONT = "Consolas"
HEADING1_SIZE = Pt(16)
TITLE_SIZE = Pt(22)
CODE_SIZE = Pt(9)


def _set_run_font(run, size=None, bold=None, mono=False):
    name = MONO_FONT if mono else FONT_NAME
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME if not mono else MONO_FONT)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size=Pt(11))
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=Pt(11))


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run_font(run, size=CODE_SIZE, mono=True)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size=Pt(10))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, size=HEADING1_SIZE if level == 1 else None, bold=True if level == 1 else None)
    return h


def add_doc_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size=TITLE_SIZE, bold=True)


def add_doc_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    _set_run_font(run, size=HEADING1_SIZE, bold=True)


def add_step(doc, step_no, title, explanation, code=None, notes=None):
    heading(doc, f"步驟 {step_no}：{title}", level=2)
    add_para(doc, explanation)
    if code:
        add_code_block(doc, code)
    if notes:
        add_para(doc, "注意事項：", bold=True)
        add_bullets(doc, notes if isinstance(notes, list) else [notes])


def section_intro(doc):
    heading(doc, "一、講義說明", level=1)
    add_para(
        doc,
        "本講義對應「manage/ 後台程式架構」教學第 3 部分之 addin.php 專題。"
        "以 manage/news/addin.php 為主範例（完整內容單元），"
        "並對照 manage/class1/addin.php（開發範本）。"
        "addin.php 是表單 POST 的唯一寫入端點，add.php 與 update.php 共用同一支。",
    )
    add_table(
        doc,
        ["檔案", "職責"],
        [
            ["add.php / update.php", "顯示表單（GET）；準備資料；require _detail.php"],
            ["_detail.php", "表單 HTML；action 指向 addin.php"],
            ["_submit.php", "hidden 欄位（PKey、manNo、csrf_token、列表搜尋條件等）"],
            ["addin.php", "驗證 → 上傳 → 寫 DB → 導回列表"],
            ["_return_list.php", "成功後組列表 URL 並 alert「新增/修改成功」"],
        ],
    )

    heading(doc, "1.1 整體資料流", level=2)
    add_bullets(
        doc,
        [
            "使用者於 _detail.php 填寫表單，按「送出」→ POST 至同目錄 addin.php",
            "_inc.php 已將 $_POST / $_GET 合併為 $filter_array；$_FILES 為 $file_array",
            "formPKey <= 0 → INSERT（新增）；formPKey > 0 → UPDATE（編輯）",
            "任一驗證失敗 → crud_form_error_redirect() 導回 add.php 或 update.php",
            "全部成功 → crud_upsert_master() + 子表 → require _return_list.php",
        ],
    )


def section_steps(doc):
    heading(doc, "二、逐步註解（manage/news/addin.php）", level=1)

    add_step(
        doc, 1, "檔頭與環境載入",
        "宣告嚴格型別；若表單含 CKEditor 須先設 $manage_csp_editor = true，"
        "讓 _inc.php 送出較寬鬆的 CSP（允許編輯器資源）。"
        "接著載入後台 Bootstrap 與模組權限。",
        """<?php
declare(strict_types=1);

$manage_csp_editor = true;          // 有富文字編輯器時設 true
require_once '../_inc.php';         // Session、DB、語系、$filter_array
require_once '../_module.php';      // manNo/subNo、權限、$Module_PKey""",
        [
            "無 CKEditor 的單元（如 weblink）可省略 $manage_csp_editor",
            "_inc.php 未完成登入檢查時會導向 login/index.php",
        ],
    )

    add_step(
        doc, 2, "讀取模組設定（_config.php）",
        "透過 manage_detail_set_config() 將 _config.php 註冊到 $GLOBALS['manage_detail_config']，"
        "再 manage_detail_tables() 取出主表與子表名稱。複製新模組時主要改 _config.php。",
        """$detailConfig = require __DIR__ . '/_config.php';
manage_detail_set_config($detailConfig);

$tables     = manage_detail_tables();
$table_name = $tables['master'];           // 例：news
$table_lang = (string)($tables['lang'] ?? '');   // news_lang
$table_msg  = (string)($tables['msg'] ?? '');    // news_msg
$table_img  = ...;                           // news_img（空則 master_img）
$FKName     = $tables['fk'];                 // 例：News_PKey
$moduleCol  = $tables['module_pk_col'];      // 預設 Module_PKey""",
        [
            "master / fk 為必填；lang、msg、img 無表可設空字串",
            "csrf 鍵名須與 _config.php 一致，且全站唯一",
        ],
    )

    add_step(
        doc, 3, "CSRF 驗證",
        "表單 hidden 的 csrf_token 必須與 Session 內 token 相符。"
        "crud_csrf_verify_form() 驗證失敗會直接終止（不消耗 token，可重送）。",
        """$csrfKey = (string)($detailConfig['csrf'] ?? 'news_addin');
crud_csrf_verify_form($csrfKey);""",
        ["列表頁刪除用 crud_csrf_guard_list()，鍵名通常為 {模組}_list"],
    )

    add_step(
        doc, 4, "取得工作上下文變數",
        "從全域輸入與 Session 取出主鍵、模組鍵、操作者，並預先計算失敗時的返回 URL。",
        """global $filter_array, $file_array, $Layer, $Class_Name;
$file_array = $file_array ?? [];

$WorkFile = (string)($_SERVER['PHP_SELF'] ?? 'addin.php');
$Login_ID = (string)($_SESSION['Login_ID'] ?? '');

$formPKey = safe_int($filter_array['PKey'] ?? 0);   // 編輯時 > 0
if ($formPKey <= 0) {
    $formPKey = safe_int($GLOBALS['Update_PKey'] ?? 0);
}

$modulePKey = (int)($GLOBALS['Module_PKey'] ?? 0);
if ($modulePKey <= 0) {
    $modulePKey = safe_int($filter_array['manNo'] ?? 0);
}

$returnUrl = crud_addin_return_url($formPKey);
// formPKey=0 → add.php?manNo=…
// formPKey>0 → update.php?PKey=…&manNo=…""",
        [
            "PKey 來自 _submit.php 的 hiddenNumeric('PKey', …)",
            "manNo 對應 module_p.PKey，寫入主檔 Module_PKey 欄",
        ],
    )

    add_step(
        doc, 5, "欄位驗證（累加 $MSG）",
        "各驗證函式回傳錯誤字串，以 .= 串接。全部通過後 $MSG 仍為空字串。"
        "news 比 class1 多了刊登日期、分類層級等驗證。",
        """$MSG = crud_addin_validate_publish_dates($filter_array);
$MSG .= crud_addin_validate_strdate($filter_array);
$MSG .= crud_validate_lang_show_strname($filter_array);
$MSG .= crud_addin_validate_layer_classes($filter_array, $Layer);

// 多語系內文（CKEditor）常以 Base64 傳入，需解碼並檢查大小
$photoSlots = crud_resolve_photo_upload_slots($filter_array, $file_array, 7);
$MAX_CONTENT_BYTES = $photoSlots['slot_max'] * 1024 * 1024;
$b64 = crud_decode_b64_content_multilang($filter_array, $MAX_CONTENT_BYTES, 6);
$MSG .= $b64['error'];
$DecodedContents = $b64['contents'];   // 解碼後 HTML，供寫入 msg 子表""",
        [
            "crud_validate_lang_show_strname：有勾「顯示」的語系，標題不可空白",
            "crud_addin_validate_layer_classes：依 module_p.intLayer 檢查 Class1～3",
            "簡化單元可只保留 strName 與 Upload 等自訂驗證（見 class1）",
        ],
    )

    add_step(
        doc, 6, "圖片上傳處理",
        "先確認 Upload 目錄可寫，再依表單槽位處理 $_FILES。"
        "成功後 $Photo/$PhotoW/$PhotoH 為各槽檔名與尺寸；$forderVal 為 Ym 子目錄。",
        """$uploadDirInfo = crud_upload_dir();
$upload_foder  = $uploadDirInfo['dir'];
$MSG          .= $uploadDirInfo['error'];

$uploadResult = crud_upload_file_slots($file_array, $upload_foder, $indices, [
    'forder_prefix' => $ForderName,      // 例：news_
    'size_bytes'    => $size_bytes,      // 預設 2MB
    'allowed_exts'  => ['gif','jpg',...],
    'field_prefix'  => 'Photo',          // Photo1, Photo2…
    'resize_thumb'  => true,
]);

$MSG .= (string)($uploadResult['messages'] ?? '');
$forderVal = $uploadResult['monthdir'] ?? date('Ym');

// PhotoM1～PhotoMn：使用者選擇的圖片版型（上圖下文等）
for ($i = 1; $i <= $maxSlots; $i++) {
    if (isset($filter_array['PhotoM' . $i])) {
        $PhotoM[$i] = (string)$filter_array['PhotoM' . $i];
    }
}""",
        [
            "未上傳新檔時，crud_save_img_slots 會保留原檔名（編輯模式）",
            "上傳錯誤訊息併入 $MSG，與欄位驗證一併處理",
        ],
    )

    add_step(
        doc, 7, "驗證失敗：導回表單",
        "只要有任何錯誤，不寫入資料庫，以 alert 顯示並 location 回 add/update。",
        """if ($MSG !== '') {
    crud_form_error_redirect($MSG, $returnUrl);
}""",
        [
            "錯誤也會寫入 $_SESSION['form_error']（若 useSession=true）",
            "使用者修正後可再次送出，CSRF token 仍有效",
        ],
    )

    add_step(
        doc, 8, "正規化日期欄位",
        "驗證通過後，將表單日期轉成 DB 儲存格式（OpenDate、EndDate、strDate 等）。",
        """crud_addin_resolve_publish_dates($filter_array);
crud_addin_resolve_strdate($filter_array);""",
        ["須在組 $data_array 之前呼叫，因會修改 $filter_array 內容"],
    )

    add_step(
        doc, 9, "編輯權限檢查（橫向越權防護）",
        "更新時確認該 PKey 的 Module_PKey 與目前 manNo 一致，防止改到其他單元資料。",
        """if ($formPKey > 0 && $modulePKey > 0) {
    $row = crud_fetch_one(
        'SELECT `Module_PKey` FROM `news` WHERE `PKey` = :pk LIMIT 1',
        ['pk' => $formPKey]
    );
    if ($row === null || (int)$row['Module_PKey'] !== $modulePKey) {
        crud_form_error_redirect('查無要修改資料或無權限', $returnUrl);
    }
}""",
        ["新增時 formPKey=0，此段跳過", "亦可改用 crud_addin_verify_master_module()"],
    )

    add_step(
        doc, 10, "組裝主檔 $data_array",
        "只放主表欄位；多語系標題在 lang 子表，不在此處。"
        "一律經 SqlFilter 過濾；crud_filter_row_for_table 移除主表不存在的欄位。",
        """$data_array = [
    'Module_PKey' => SqlFilter($modulePKey, 'int'),
    'strName'     => SqlFilter((string)($filter_array['strName1'] ?? ''), 'tab'),
    'dtUDate'     => date('Y-m-d H:i:s'),
    'UserID'      => SqlFilter($Login_ID, 'tab'),
];
if (isset($filter_array['Upload'])) {
    $data_array['Upload'] = SqlFilter($filter_array['Upload'], 'tab');
}
$data_array = crud_addin_append_publish_master_fields(
    $data_array, $filter_array, $Layer, $table_name
);
// 依表結構追加 Keywords、Description、Class1～3、OpenDate…

$data_array = crud_filter_row_for_table($table_name, $data_array);""",
        [
            "strName1 為第一語系標題；主表 strName 常存預設語系",
            "class1 範本另含 Sort、Home 等分類專用欄位",
            "自訂欄位加在 crud_filter_row_for_table 之前",
        ],
    )

    add_step(
        doc, 11, "寫入資料庫（try 區塊）",
        "crud_upsert_master 依 formPKey 決定 INSERT 或 UPDATE，回傳新/舊主鍵 parentPKey。"
        "再依序寫入 lang、msg、img 子表；有標籤則 tag_relation_save。",
        """try {
    $upsert     = crud_upsert_master($table_name, $formPKey, $data_array);
    $parentPKey = $upsert['pkey'];
    $show       = $upsert['action'];    // 「新增成功!」或「修改成功!」

    if ($table_lang !== '') {
        crud_save_lang_slots($table_lang, $FKName, $parentPKey, $filter_array);
    }
    if ($table_msg !== '') {
        crud_save_msg_blocks_multilang(
            $table_msg, $FKName, $parentPKey,
            $DecodedContents, $filter_array, 6
        );
    }
    if ($table_img !== '') {
        crud_save_img_slots(
            $table_img, $FKName, $parentPKey, $forderVal,
            $Photo, $PhotoW, $PhotoH, $PhotoM,
            $filter_array, $maxSlots, $upload_foder
        );
    }

  // 標籤（須 _inc.php 開啟 tag 且 _config 有 tag_relation_parent_col）
    if (manage_module_show_detail_field('tag')) {
        tag_relation_save($parentPKey, 'News_PKey', $filter_array);
    }

    $actionShow = $show;
    require_once '../_return_list.php';
    exit;""",
        [
            "子表寫入失敗會拋例外，進入 catch",
            "manage_history 在 crud_upsert_master 內自動記錄",
            "lang 表含 strName、Title、Description、Keywords 等（依欄位存在與否）",
        ],
    )

    add_step(
        doc, 12, "例外處理（catch）",
        "DB 或檔案錯誤時寫 sql_error log，Session 設 error_msg，導回表單。",
        """} catch (Throwable $e) {
    sql_error('', $e->getMessage(), $WorkFile, $Login_ID, ...);
    $_SESSION['error_msg'] = '資料寫入失敗';
    header('Location: ' . crud_addin_return_url($formPKey));
    exit;
}""",
        ["正式環境勿將 $e->getMessage() 直接顯示給使用者"],
    )

    add_step(
        doc, 13, "成功導回（_return_list.php）",
        "addin 設定 $actionShow 後引入 _return_list.php："
        "從 POST 的 Q_* hidden 還原列表搜尋條件，alert 成功訊息後導向 list.php。",
        """// _return_list.php 摘要：
$params = ['Send' => '搜尋'];
// 還原 Page、PageSize、Q_Keywords、Q_Class1…、manNo、subNo
$url = $list . '?' . http_build_query($params);
manage_alert_script($actionShow, $url);  // alert + location""",
        [
            "list、list_module 來自 _submit.php hidden 欄位",
            "按「關閉」不經 addin，由 JS 直接 POST 至 _return_list.php",
        ],
    )


def section_compare(doc):
    heading(doc, "三、news 與 class1 範本差異", level=1)
    add_table(
        doc,
        ["項目", "news/addin.php", "class1/addin.php"],
        [
            ["$manage_csp_editor", "有（多語系 CKEditor）", "無"],
            ["刊登日期驗證", "crud_addin_validate_publish_dates 等", "無"],
            ["分類 Layer 驗證", "crud_addin_validate_layer_classes", "無"],
            ["Sort 驗證", "無（has_sort=false）", "Sort < 0 檢查"],
            ["主檔欄位", "OpenDate、Class、Keywords…", "Sort、Home、Upload"],
            ["圖片槽位數", "最多 7（內容圖）", "1（列表圖）"],
            ["標籤關聯", "tag_relation_save", "無"],
            ["註解定位", "結構同 class1", "官方複製範本"],
        ],
    )


def section_customize(doc):
    heading(doc, "四、複製到新模組時的修改清單", level=1)
    add_table(
        doc,
        ["順序", "修改處", "說明"],
        [
            ["1", "_config.php", "master、子表名、fk、csrf（唯一）"],
            ["2", "csrfKey 預設字串", "與 _config csrf 一致"],
            ["3", "驗證區", "刪除不需要的 validate_*；保留 strName、Upload"],
            ["4", "圖片區", "調整 maxSlots、ForderName、allowed_exts"],
            ["5", "主檔 $data_array", "對齊資料表欄位；用 crud_table_has_column 防呆"],
            ["6", "子表寫入", "無 msg 表則移除 crud_save_msg_blocks 區塊"],
            ["7", "標籤", "無標籤則移除 tag_relation 區塊"],
            ["8", "測試", "新增、編輯、驗證失敗、越權 PKey、圖片上傳"],
        ],
    )


def section_form_fields(doc):
    heading(doc, "五、表單欄位與 $filter_array 對照（news）", level=1)
    add_table(
        doc,
        ["表單 name", "用途", "寫入位置"],
        [
            ["PKey", "編輯主鍵", "決定 UPDATE；不直接寫入 data_array"],
            ["manNo / subNo", "模組導覽", "modulePKey → Module_PKey"],
            ["csrf_token", "CSRF", "crud_csrf_verify_form"],
            ["strName1～strName6", "各語系標題", "news_lang.strName"],
            ["Contents1_b64～", "各語系 HTML 內文", "news_msg（解碼後）"],
            ["Photo1～Photo7", "圖片檔案", "news_img"],
            ["PhotoM1～", "圖片版型", "news_img.PhotoM"],
            ["Upload", "上下架", "news.Upload"],
            ["OpenDate / EndDate", "刊登區間", "news 主表"],
            ["Class1～Class3", "分類", "news 主表（依 Layer）"],
            ["Q_Keywords 等 hidden", "列表搜尋條件", "僅 _return_list 使用，不寫 DB"],
        ],
    )


def section_faq(doc):
    heading(doc, "六、常見問題", level=1)
    add_table(
        doc,
        ["現象", "可能原因", "處理方式"],
        [
            ["CSRF 驗證失敗", "csrf key 與表單不一致；Session 過期", "檢查 _config csrf 與 add.php crud_csrf_ensure_page 是否同 key"],
            ["送出後空白或 500", "manage_detail_set_config 未呼叫；表名錯", "確認 require _config 與 master 表存在"],
            ["英文語系沒存到", "crud_save_lang_slots 未執行；lang 表無資料", "確認 _config lang 鍵與表單 strName{n}"],
            ["圖片有選但沒存", "超過 size_bytes；目錄權限", "查 $MSG 與 Upload 目錄 is_writable"],
            ["修改成功但列表看不到", "Module_PKey 不符；Upload=No", "查主檔 Module_PKey 與列表 WHERE"],
            ["查無要修改資料或無權限", "PKey 屬於其他 manNo", "正常防護；檢查 URL manNo 與資料歸屬"],
        ],
    )


def section_exercise(doc):
    heading(doc, "七、實作練習（講師用）", level=1)
    add_bullets(
        doc,
        [
            "學員複製 news/addin.php，刪除 msg 與標籤區塊，改為僅標題+上下架的最小版本",
            "在驗證區新增：Upload=Yes 時 OpenDate 必填",
            "使用 Xdebug 於 crud_upsert_master 設斷點，觀察 INSERT 與 UPDATE 分支",
            "故意送錯 csrf_token，確認 crud_csrf_verify_form 行為",
            "對照 _return_list.php，確認 Q_Keywords 是否正確回到列表搜尋",
        ],
    )


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_doc_title(doc, "addin.php 逐步註解講義")
    add_doc_subtitle(doc, "後台教學第 3 部分｜以 manage/news/addin.php 為例")

    section_intro(doc)
    section_steps(doc)
    section_compare(doc)
    section_customize(doc)
    section_form_fields(doc)
    section_faq(doc)
    section_exercise(doc)

    try:
        doc.save(str(OUT))
        print("Saved:", OUT)
    except PermissionError:
        doc.save(str(OUT_ALT))
        print("Saved (alt):", OUT_ALT)


if __name__ == "__main__":
    main()
