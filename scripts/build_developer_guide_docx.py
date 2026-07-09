#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""產生 文件/新進開發人員上手指南.docx"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "文件" / "新進開發人員上手指南.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("brick6 新進開發人員上手指南", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_para(
        doc,
        "本文件協助第一次接觸 brick6 的 PHP 開發人員，快速理解前後端架構、日常維運，"
        "以及如何進行局部功能修改（後台欄位增減、複製模組、前台調整）。"
        "第一週實戰練習（Day 6/7）另見 docs/ONBOARDING-WEEK1.md 與 文件/ONBOARDING-WEEK1.docx。",
    )

    # ── 1. 專案概觀 ──
    add_heading(doc, "1. 專案概觀", 1)
    add_para(doc, "brick6 是 PHP 8.4 單體式 CMS，採「根目錄 PHP + include + 模組化 CRUD」架構，無 Laravel/Symfony 框架。")
    add_table(
        doc,
        ["區塊", "位置", "說明"],
        [
            ["前台", "根目錄 *.php", "訪客網站"],
            ["後台", "manage/{module}/", "內容 CRUD"],
            ["共用", "include/", "DB、CRUD、前台 helper"],
            ["設定", ".env", "不可提交 Git"],
            ["路由", ".htaccess", "news.htm → news.php"],
        ],
    )
    add_para(doc, "核心原則：manage/{module}/_config.php 為 schema 唯一來源；前台 require 合併同一份設定。")

    # ── 2. 目錄結構 ──
    add_heading(doc, "2. 目錄結構速查", 1)
    add_bullets(
        doc,
        [
            "_inc.php — 前台 bootstrap",
            "manage/_inc.php — 後台 bootstrap",
            "manage/{module}/_config.php — 模組設定（最重要）",
            "manage/{module}/list.php、add.php、update.php、addin.php — CRUD 四件套",
            "include/crud_helpers.php — 後台共用 CRUD",
            "include/frontend_helpers.php — 前台查詢/SEO",
            "sql/ — migration（手動執行）",
            "Upload/ — 上傳檔案（不進 Git）",
        ],
    )

    add_heading(doc, "後台四件套職責", 2)
    add_table(
        doc,
        ["檔案", "職責"],
        [
            ["list.php", "列表、搜尋、刪除、排序"],
            ["add.php / update.php", "準備表單（不寫 DB）"],
            ["_detail.php", "表單 HTML"],
            ["addin.php", "POST 驗證、寫 DB、上傳"],
        ],
    )

    # ── 3. 請求流程 ──
    add_heading(doc, "3. 請求流程", 1)
    add_heading(doc, "前台", 2)
    add_para(doc, "Browser → .htaccess → {module}.php → _inc.php → frontend_module_set_config → 查詢 → _header + 內容 + _footer")
    add_heading(doc, "後台", 2)
    add_para(doc, "list.php → add/update.php → _detail.php → addin.php（寫 DB）→ redirect list.php")
    add_heading(doc, "AJAX", 2)
    add_bullets(doc, ["frontend-visit-log.php", "manage/generate_*.php", "manage/ajax/、模組內端點"])

    # ── 4. 資料庫 ──
    add_heading(doc, "4. 資料庫與模組資料表模式", 1)
    add_bullets(
        doc,
        [
            "host.php 載入 .env；sql_conn() 建立 PDO",
            "查詢一律 Prepared Statements",
            "典型：master + _lang + _msg + _img + _link + view_{module}",
            "FK 命名：{Module}_PKey（如 News_PKey）",
            "上傳目錄：{forder_prefix}{PKey}/",
        ],
    )
    add_heading(doc, "模組複雜度參考", 2)
    add_table(
        doc,
        ["複雜度", "模組", "特點"],
        [
            ["簡單", "faq", "排序、單一內文、無 link"],
            ["標準", "news / paper / company", "完整四子表"],
            ["子模組", "album_d", "依附 album"],
            ["系統", "control / webset", "帳號、網站設定"],
        ],
    )

    # ── 5. 欄位增減 ──
    add_heading(doc, "5. 局部修改：後台欄位增減", 1)
    add_numbered(
        doc,
        [
            "SQL migration（sql/*.sql），各環境手動執行",
            "_form_data.php：class1_detail_init_defaults()、class1_detail_load()",
            "_detail.php：Bootstrap 表單欄位",
            "addin.php：$data_array 或 crud_save_lang_slots()",
            "list.php（可選）：列表顯示",
            "前台頁：crud_row_val() + e() 輸出",
        ],
    )
    add_para(doc, "語系子表新欄位：可能需擴充 crud_helpers.php（參考 strNote 實作）。")

    # ── 6. 複製模組 ──
    add_heading(doc, "6. 複製模組", 1)
    add_heading(doc, "複製整個模組（10 步）", 2)
    add_numbered(
        doc,
        [
            "複製 manage/{source}/ → manage/{new}/",
            "修改 _config.php（表名、FK、csrf、forder_prefix）",
            "全域搜尋替換模組名",
            "建立 DB 表 + View；子表 FK 欄位改名",
            "module_p / module_lang 註冊選單（PageLink）",
            "複製前台 PHP 頁",
            "frontend_module_pkey_for_page('xxx.htm')",
            ".htaccess RewriteRule",
            "_code_lang.php 多語系",
            "view_* 加入 crud_helpers 白名單",
        ],
    )
    add_para(doc, "複製單筆：列表「複製」→ add.php 帶 PKey，清空 PKey/圖片以新增送出。")
    add_para(doc, "dev 練習：Day 7 faqdemo — 見 ONBOARDING-WEEK1.md。")

    # ── 7. 前台 ──
    add_heading(doc, "7. 前台架構與修改", 1)
    add_heading(doc, "頁面 Shell", 2)
    add_para(doc, "_inc.php → _in_code_head.php → _in_javascript.php → _header → _banner → main → _footer → _in_code_bottom.php")
    add_heading(doc, "修改類型", 2)
    add_table(
        doc,
        ["需求", "修改位置"],
        [
            ["後台新欄位", "列表/內頁 PHP + e()"],
            ["靜態文字", "_code_lang.php"],
            ["全站設定", "webset → _footer.php"],
            ["樣式", "css/style.css"],
            ["互動", "js/{page}-page.js"],
            ["AJAX", "根目錄端點 + $.ajax().done().fail()"],
        ],
    )
    add_heading(doc, "常用 Helper", 2)
    add_bullets(
        doc,
        [
            "frontend_module_set_config() / frontend_module_config()",
            "frontend_list_where() / frontend_fetch_list()",
            "frontend_fetch_detail() / frontend_fetch_msg_contents()",
            "frontend_detail_href() / e() / e_attr()",
        ],
    )

    # ── 8. 開發規範 ──
    add_heading(doc, "8. 開發規範", 1)
    add_bullets(
        doc,
        [
            "敏感資訊只讀 .env，禁止寫死",
            "SQL：PDO Prepared Statements",
            "XSS：e() / htmlspecialchars",
            "jQuery：$() 語法；AJAX：$.ajax().done().fail()",
            "PHP：頂部邏輯/POST，HTML 區只輸出",
            "Git：不提交 .env、Upload/",
        ],
    )

    # ── 9. 維運 ──
    add_heading(doc, "9. 維運與部署", 1)
    add_bullets(
        doc,
        [
            "本機：.env.example → .env；後台 manage/login/index.php",
            "正式機：config/env.path.php 可指向 private/.env",
            "sql/*.sql 提交後各環境手動執行",
            "選單來源：module_p、module_d、view_module_lang",
        ],
    )

    # ── 10. Checklist ──
    add_heading(doc, "10. 實務 Checklist", 1)
    add_heading(doc, "新增欄位", 2)
    add_bullets(doc, ["SQL", "_form_data.php", "_detail.php", "addin.php", "list（可選）", "前台", "測試"])
    add_heading(doc, "複製模組", 2)
    add_bullets(doc, ["manage 目錄 + _config", "DB + View + FK", "選單", "前台 + htaccess", "白名單", "測試"])

    # ── 11. 學習路徑 ──
    add_heading(doc, "11. 建議學習路徑（第一週）", 1)
    add_para(doc, "每日 3～4 小時。Day 6/7 含可執行 patch。")

    days = [
        ("Day 1：Bootstrap", ["host.php、Conn.php、_inc.php、$filter_array", "產出：請求流程圖"]),
        ("Day 2：後台 CRUD", ["faq 模組 list→addin 全流程", "產出：CRUD 流程圖"]),
        ("Day 3：前台", ["faq.php 三段式、.htaccess", "產出：Shell 結構圖"]),
        ("Day 4：crud_helpers", ["crud_cfg、crud_save_lang_slots 等", "產出：函式速查表"]),
        ("Day 5：小改動", ["order_by、_code_lang、對照 news", "產出：2 個改動驗證"]),
        ("Day 6：strNote", ["day6_faq_add_strnote.sql", "後台+前台驗收"]),
        ("Day 7：faqdemo", ["day7 SQL + onboarding_install_view_faqdemo.php", "faqdemo.htm 驗收"]),
    ]
    for day_title, items in days:
        add_heading(doc, day_title, 2)
        add_bullets(doc, items)

    # ── 12. 模組表 ──
    add_heading(doc, "12. 主要模組對照表", 1)
    add_table(
        doc,
        ["後台", "前台", "備註"],
        [
            ["manage/news/", "news.php / news-detail.php", "標準範本"],
            ["manage/faq/", "faq.php", "簡化無內頁"],
            ["manage/faqdemo/", "faqdemo.php", "Day 7 練習"],
            ["manage/product/", "product.php", "多圖 Tab"],
            ["manage/album/", "album.php", "含 album_d"],
            ["manage/webset/", "—", "全站設定"],
        ],
    )

    add_heading(doc, "相關文件", 1)
    add_bullets(
        doc,
        [
            "docs/DEVELOPER-ONBOARDING-GUIDE.md — 本文件 Markdown",
            "docs/ONBOARDING-WEEK1.md — 第一週實戰 patch",
            "文件/新進開發人員上手指南.docx — 本文件",
            "文件/ONBOARDING-WEEK1.docx — 第一週實戰",
        ],
    )

    add_para(doc, "文件版本：2026-07-09 · brick6 新進開發人員上手指南", bold=True)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
