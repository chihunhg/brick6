#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""產生 文件/ONBOARDING-WEEK1.docx"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "文件" / "ONBOARDING-WEEK1.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # 頁面邊界
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("brick6 新進開發人員 — 第一週上手指南", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_para(
        doc,
        "本文件協助第一次接觸 brick6 的 PHP 開發人員，快速理解前後端架構、維運方式，"
        "以及局部功能修改（欄位增減、模組複製）。Day 6 / Day 7 含可於 dev 環境直接執行的 SQL 與程式 patch。",
    )

    add_heading(doc, "1. 專案概觀", 1)
    add_para(doc, "brick6 是 PHP 8.4 單體式 CMS，採「根目錄 PHP + include + 模組化 CRUD」架構。")
    add_bullets(
        doc,
        [
            "前台：根目錄 *.php",
            "後台：manage/{module}/",
            "共用：include/（crud_helpers、frontend_helpers）",
            "設定：.env（不可提交 Git）",
            "路由：.htaccess 友好 URL",
        ],
    )
    add_para(doc, "核心原則：manage/{module}/_config.php 為 schema 唯一來源；前台 require 合併同一份設定。")

    add_heading(doc, "2. 建議學習路徑 Day 1～5", 1)

    days = [
        (
            "Day 1：理解 Bootstrap",
            [
                "必讀：.env.example、include/host.php、Conn.php、_inc.php、manage/_inc.php、Function.php",
                "搞懂 $filter_array、$this_lang、$Array_MU_*、$Module_PKey",
                "練習：開啟 faq.htm，對照 _inc.php 載入順序",
                "產出：請求啟動流程圖 + 全域變數對照表",
            ],
        ),
        (
            "Day 2：後台 CRUD（faq）",
            [
                "走一遍：list → add → addin → update → 刪除",
                "對照：_config.php、_form_data.php、_detail.php、addin.php",
                "理解四件套：list 看資料、add/update 準備表單、addin 寫 DB",
                "練習：新增 [TEST-Day2] 資料，phpMyAdmin 查四表",
                "產出：FAQ CRUD 流程圖",
            ],
        ),
        (
            "Day 3：前台 faq.php",
            [
                "三段式：require _inc → frontend_module_set_config → HTML 輸出",
                "對照 manage/faq/_config.php 與 faq.php 覆寫欄位",
                "讀 .htaccess：faq.htm → faq.php",
                "練習：改 order_by 驗證列表順序",
                "產出：前台 Shell 結構圖",
            ],
        ),
        (
            "Day 4：crud_helpers.php",
            [
                "掌握：crud_cfg、crud_module_where、crud_process_list_actions",
                "掌握：crud_save_lang_slots、crud_fetch_all、delete_cascade_by_ids",
                "從 list.php / addin.php Ctrl+Click 追函式",
                "產出：常用 CRUD 函式速查表（10～15 個）",
            ],
        ),
        (
            "Day 5：frontend_helpers 小改動",
            [
                "函式：frontend_list_where、frontend_fetch_faq_items、frontend_module_set_config",
                "練習 A：改 faq.php 的 order_by",
                "練習 B：改 _code_lang.php 靜態文字",
                "練習 C：對照 news.php 分頁與側欄",
                "產出：完成 2 個小改動並記錄驗證方式",
            ],
        ),
    ]
    for day_title, items in days:
        add_heading(doc, day_title, 2)
        add_numbered(doc, items)

    add_heading(doc, "3. Day 6 實戰：FAQ 新增 strNote 欄位", 1)
    add_para(doc, "目標：在 faq_lang 新增備註 strNote，完成後台表單 → 寫入 → 前台顯示。")

    add_heading(doc, "3.1 執行 SQL", 2)
    add_para(doc, "mysql -u USER -p DB_NAME < sql/onboarding/day6_faq_add_strnote.sql")

    add_heading(doc, "3.2 程式 patch（已納版控）", 2)
    add_bullets(
        doc,
        [
            "include/crud_helpers.php — strNote 語系讀寫",
            "manage/faq/_form_data.php — 預設值與載入",
            "manage/faq/_detail.php — 表單欄位 strNote{n}",
            "include/frontend_helpers.php — frontend_fetch_faq_items 回傳 note",
            "faq.php — 顯示 .faqItem__note",
            "css/style.css — 備註樣式",
        ],
    )

    add_heading(doc, "3.3 驗收", 2)
    add_numbered(
        doc,
        [
            "後台 FAQ 新增，備註填 [Day6-test]",
            "faq_lang.strNote 有值",
            "前台 faq.htm 答案上方顯示備註",
            "編輯頁可帶出備註",
        ],
    )

    add_heading(doc, "4. Day 7 實戰：複製 FAQ 為 faqdemo", 1)
    add_para(doc, "目標：獨立模組 faqdemo，含後台 CRUD、前台 faqdemo.htm，與 faq 資料隔離。")

    add_heading(doc, "4.1 執行 SQL", 2)
    add_para(doc, "mysql -u USER -p DB_NAME < sql/onboarding/day7_faqdemo_module.sql")

    add_heading(doc, "4.2 建立 view_faqdemo", 2)
    add_para(doc, "php scripts/onboarding_install_view_faqdemo.php")
    add_para(doc, "或 SHOW CREATE VIEW view_faq 後手動替換 faq → faqdemo、FAQ_PKey → FAQDemo_PKey。")

    add_heading(doc, "4.3 程式 patch（已納版控）", 2)
    add_bullets(
        doc,
        [
            "manage/faqdemo/ — 完整 CRUD 目錄",
            "faqdemo.php — 前台頁",
            "js/faqdemo-page.js",
            ".htaccess — faqdemo.htm 規則",
            "include/crud_helpers.php — view_faqdemo 白名單",
        ],
    )

    add_heading(doc, "4.4 驗收", 2)
    add_numbered(
        doc,
        [
            "後台 FAQ Demo 列表可開啟",
            "新增資料，Upload/faqdemo_{PKey}/ 可上傳",
            "前台 faqdemo.htm 顯示資料",
            "原 faq.htm 不受影響",
        ],
    )

    add_heading(doc, "5. 欄位增減 SOP", 1)
    add_numbered(
        doc,
        [
            "SQL migration（sql/*.sql）",
            "_form_data.php：init + load",
            "_detail.php：表單 UI",
            "addin.php 或 crud_save_lang_slots 寫入",
            "list.php（可選）",
            "前台頁 + helper",
            "端到端測試",
        ],
    )

    add_heading(doc, "6. 複製模組 Checklist", 1)
    add_numbered(
        doc,
        [
            "複製 manage/{module}/，改 _config.php",
            "SQL 建表 + View + FK 欄位名",
            "module_p / module_lang 選單",
            "前台 {module}.php + .htaccess",
            "view_* 白名單",
            "端到端測試",
        ],
    )

    add_heading(doc, "7. 開發規範", 1)
    add_bullets(
        doc,
        [
            "敏感資訊只讀 .env",
            "SQL 一律 Prepared Statements",
            "輸出用 e() / e_attr() 防 XSS",
            "jQuery：$.ajax().done().fail()",
            "不提交 .env、Upload/",
        ],
    )

    add_heading(doc, "8. 變更檔案索引", 1)
    files = [
        "sql/onboarding/day6_faq_add_strnote.sql",
        "sql/onboarding/day7_faqdemo_module.sql",
        "scripts/onboarding_install_view_faqdemo.php",
        "include/crud_helpers.php",
        "include/frontend_helpers.php",
        "manage/faq/_form_data.php",
        "manage/faq/_detail.php",
        "manage/faqdemo/",
        "faq.php",
        "faqdemo.php",
        "docs/ONBOARDING-WEEK1.md",
    ]
    add_bullets(doc, files)

    add_para(doc, "文件版本：2026-07-09 · brick6 onboarding week 1", bold=True)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
