# -*- coding: utf-8 -*-
"""Generate Word doc: 後台新增／調整單元指南"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parent / (
    "\u5f8c\u53f0\u65b0\u589e\u8207\u8abf\u6574\u55ae\u5143"
    "-\u4fee\u6539\u7a0b\u5f0f\u6e05\u55ae\u8207\u6ce8\u610f\u4e8b\u9805.docx"
)
OUT_ALT = Path(__file__).resolve().parent / (
    "\u5f8c\u53f0\u65b0\u589e\u8207\u8abf\u6574\u55ae\u5143"
    "-\u4fee\u6539\u7a0b\u5f0f\u6e05\u55ae\u8207\u6ce8\u610f\u4e8b\u9805-\u66f4\u65b0.docx"
)

FONT_NAME = "Microsoft JhengHei"
HEADING1_SIZE = Pt(16)
TITLE_SIZE = Pt(22)


def _set_run_font(run, size=None, bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size=Pt(11))
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=Pt(11))


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, size=HEADING1_SIZE if level == 1 else None, bold=True if level == 1 else None)
    return h


def add_doc_title(doc, text):
    """主標：置中、字級較大"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size=TITLE_SIZE, bold=True)
    return p


def add_doc_subtitle(doc, text):
    """副標：置中、字級同第一層章節標題"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    _set_run_font(run, size=HEADING1_SIZE, bold=True)
    return p


def section_preface(doc):
    heading(doc, "\u5efa\u8b70\u8907\u88fd\u7bc4\u672c", level=1)
    add_para(
        doc,
        "\u65b0\u589e\u5f8c\u53f0\u5167\u5bb9\u55ae\u5143\u6642\uff0c\u8acb\u5148\u4f9d\u55ae\u5143\u985e\u578b\u9078\u64c7\u6700\u63a5\u8fd1\u7684\u73fe\u6709 manage \u76ee\u9304\u8907\u88fd\uff0c"
        "\u518d\u6309\u672c\u6587\u4ef6\u5404\u7ae0\u7bc0\u4fee\u6539\u8cc7\u6599\u8868\u3001\u6a94\u6848\u8207\u5171\u7528\u8a2d\u5b9a\u3002"
        "\u8907\u88fd\u5b8c\u6210\u5f8c\uff0c\u9805\u76ee\u5305\u542b list\uff08\u7d22\u5f15\uff09\u3001add\uff0fupdate\uff08\u8868\u55ae\uff09\u8207 addin\uff08\u5132\u5b58\uff09\uff1b"
        "\u5b50\u6a21\u7d44\uff08\u5982\u76f8\u518a\u5716\u7247\u3001\u554f\u5377\u985e\u5225\uff09\u5247\u984f\u7528 manage_child_* \u5171\u7528\u9aa8\u67b6\u3002"
        "\u4e0b\u8868\u70ba\u5404\u985e\u578b\u55ae\u5143\u7684\u5efa\u8b70\u8907\u88fd\u4f86\u6e90\u3002",
    )
    add_table(
        doc,
        ["\u55ae\u5143\u985e\u578b", "\u5efa\u8b70\u8907\u88fd\u4f86\u6e90", "\u8aaa\u660e"],
        [
            ["\u4e00\u822c\u5167\u5bb9\uff08\u5206\u985e\uff0b\u591a\u8a9e\u7cfb\uff0b\u5716\uff0b\u5167\u6587\uff09", "manage/paper/ \u6216 manage/news/", "\u7d50\u69cb\u6700\u5b8c\u6574"],
            ["\u7c21\u5316\u5167\u5bb9\uff08\u5c11\u6b04\u4f4d\uff09", "manage/weblink/\u3001manage/faq/", "\u7121 msg\uff0f\u6b04\u4f4d\u8f03\u5c11"],
            ["\u5b50\u6a21\u7d44\uff08\u4f9d\u9644\u4e3b\u55ae\u5143\uff09", "manage/album_d/\u3001manage/question_class/", "\u8d70 manage_child_* \u9aa8\u67b6"],
            ["\u55ae\u5143\u8a3b\u518a\u672c\u8eab", "\u4e0d\u8907\u88fd\uff0c\u7528\u5f8c\u53f0\u300c\u55ae\u5143\u8a2d\u5b9a\u300d", "\u5c0d\u61c9 manage/module/"],
        ],
    )


def section_1(doc):
    heading(doc, "\u4e00\u3001\u65b0\u589e\u55ae\u5143\uff08\u8907\u88fd\uff0b\u8a3b\u518a\uff09", level=1)

    heading(doc, "1. \u8cc7\u6599\u5eab\uff08\u5148\u65bc\u7a0b\u5f0f\uff09", level=2)
    add_table(
        doc,
        ["\u9805\u76ee", "\u8aaa\u660e"],
        [
            ["\u4e3b\u8868 {name}", "\u542b Module_PKey\u3001Sort\u3001Upload \u7b49"],
            ["\u5b50\u8868", "\u4f9d\u9700\u6c42\uff1a{name}_img\u3001{name}_lang\u3001{name}_msg\u3001{name}_link\uff0f{name}_relation"],
            ["program", "\u65b0\u7a0b\u5f0f\u985e\u578b\u6642\u65b0\u589e\uff08strLink\uff1d\u76ee\u9304\u540d\u3001MaxLayer\u3001isList\u3001isDetail\u3001isColum\uff09"],
            ["module_p\uff0fmodule_d", "\u5f8c\u53f0\u300c\u55ae\u5143\u8a2d\u5b9a\u300d\u5efa\u7acb\uff0c\u6216\u624b\u5de5 INSERT"],
        ],
    )

    heading(doc, "2. \u8907\u88fd\u76ee\u9304", level=2)
    add_para(doc, "manage/{範本}/  \u2192  manage/{新單元}/")
    add_para(doc, "\u4e00\u822c\u5167\u5bb9\u55ae\u5143\u5178\u578b\u6a94\u6848\uff0810 \u500b\uff09\uff1a", bold=True)
    add_bullets(
        doc,
        [
            "manage/{新單元}/_config.php",
            "manage/{新單元}/list.php",
            "manage/{新單元}/_list.php",
            "manage/{新單元}/add.php",
            "manage/{新單元}/update.php",
            "manage/{新單元}/addin.php",
            "manage/{新單元}/_detail.php",
            "manage/{新單元}/_form_data.php\uff08\u6709\u81ea\u8a02 load/save \u624d\u9700\u8981\uff09",
            "manage/{新單元}/_upload.php",
            "manage/{新單元}/_del_img.php",
        ],
    )

    heading(doc, "3. \u5fc5\u6539\u6a94\u6848\uff08\u65b0\u55ae\u5143\u76ee\u9304\u5167\uff09", level=2)
    add_table(
        doc,
        ["\u6a94\u6848", "\u4fee\u6539\u91cd\u9ede"],
        [
            ["_config.php", "master\u3001img/lang/msg/link\u3001fk\u3001csrf\u3001list_csrf\u3001forder_prefix\u3001photo_slots\u3001tag_relation_parent_col \u7b49"],
            ["list.php", "table_name\u3001FKName\u3001list_csrf\uff1bcrud_cfg()\uff1b\u641c\u5c0b\u6b04\uff08\u8981\u6c42\u8981\u4e00\u81f4\u7b2c\u56db\u7ae0\uff09"],
            ["_list.php", "\u5217\u8868\u6b04\u4f4d\u3001tableGrid class\u3001\u7de8\u8f2f\u9023\u7d50"],
            ["add.php\uff0fupdate.php", "\u8f09\u5165 _detail.php\uff1b\u8907\u88fd\u4f86\u6e90 PKey\uff08\u82e5\u6709\uff09\uff08\u8981\u6c42\u8981\u4e00\u81f4\u7b2c\u4e94\u7ae0\uff09"],
            ["_detail.php", "\u8868\u55ae HTML\u3001JS fieldCheck*\u3001\u5716\u7247\u69fd\u4f4d\u3001\u6b04\u4f4d\u986f\u793a\u689d\u4ef6"],
            ["addin.php", "\u9a57\u8b49\u3001\u5beb\u5165\u4e3b\u8868\uff0f\u5b50\u8868\uff0f\u4e0a\u50b3\uff1bcrud_csrf_verify_form \u7684 key"],
            ["_form_data.php", "form bag\u3001form_load\uff0fform_save\u3001\u8907\u88fd\u908f\u8f2f"],
            ["_upload.php", "\u901a\u5e38\u50c5 require '../_upload_endpoint.php'"],
            ["_del_img.php", "\u522a\u5716 API\uff0c\u78ba\u8a8d $detailConfig \u4e00\u81f4"],
        ],
    )

    heading(doc, "4. \u5171\u7528\u6a94\u6848\uff08\u4f9d\u9700\u6c42\u4fee\u6539\uff09", level=2)
    add_table(
        doc,
        ["\u6a94\u6848", "\u6642\u6a5f"],
        [
            ["manage/_inc.php", "\u65b0\u55ae\u5143\u8981\u958b\u95dc\u9996\u9801\uff0f\u7c21\u8ff0\uff0f\u5217\u8868\u5716\uff0f\u6a19\u7c64 \u2192 $manage_module_detail_fields\uff08\u8981\u6c42\u8981\u4e00\u81f4\u7b2c\u516d\u7ae0\uff09"],
            ["manage/css/style.css", "_list.php \u65b0\u589e\u6b04\u4f4d \u2192 .tableGrid--{單元名}"],
            ["manage/class1/_config.php\uff08\u53ca class2\uff0f3\uff09", "\u65b0\u4e3b\u8868\u6709 Class1_PKey \u2192 delete_lock_tables"],
            ["\u7236\u55ae\u5143 _list.php", "\u5b50\u6a21\u7d44\u5165\u53e3\u6309\u9215"],
        ],
    )

    heading(doc, "5. \u5f8c\u53f0\u8a3b\u518a\uff08manage/module/\uff09", level=2)
    add_para(doc, "\u5f8c\u53f0\u300c\u55ae\u5143\u8a2d\u5b9a\u300d\u65b0\u589e module_p\uff1a\u540d\u7a31\u3001intUse\uff08\u5c0d\u61c9 program\uff09\u3001intLayer\uff1b\u5132\u5b58\u5f8c\u7522\u751f module_d \u5074\u6b04\u3002")


def section_2(doc):
    heading(doc, "\u4e8c\u3001\u65e2\u6709\u55ae\u5143\u300c\u8abf\u6574\u6b04\u4f4d\u300d", level=1)

    heading(doc, "A. \u5171\u7528\u6b04\u4f4d\uff08home\uff0finterview\uff0flist\uff0ftag\u2026\uff09", level=2)
    add_para(doc, "\u8a73\u7d30\u8a2d\u5b9a\u8981\u6c42\u8981\u4e00\u81f4\u7b2c\u516d\u7ae0\u3002")
    add_table(
        doc,
        ["\u6a94\u6848", "\u52d5\u4f5c"],
        [
            ["manage/_inc.php", "$manage_module_detail_fields[Module_PKey] \u8a2d true\uff0ffalse"],
            ["manage/{單元}/_detail.php", "manage_module_show_detail_field('xxx') \u5305\u4f4f\u5340\u584a"],
            ["manage/{單元}/_list.php", "$showHomeColumn = manage_module_show_detail_field('home') \u7b49"],
            ["manage/{單元}/addin.php", "\u5132\u5b58\u908f\u8f2f\u4f9d $showXxxField \u6c7a\u5b9a\u662f\u5426\u5beb\u5165"],
        ],
    )

    heading(doc, "B. \u81ea\u8a02\u6b04\u4f4d\uff08\u589e\uff0f\u522a\uff0f\u6539\uff09", level=2)
    add_table(
        doc,
        ["\u6a94\u6848", "\u65b0\u589e\u6b04\u4f4d", "\u522a\u9664\u6b04\u4f4d"],
        [
            ["_detail.php", "\u52a0 input\uff0bfieldCheck", "\u522a HTML\uff0bJS \u9a57\u8b49"],
            ["addin.php", "\u52a0\u9a57\u8b49\u3001$row \u5beb\u5165", "\u522a\u9a57\u8b49\u8207 assign"],
            ["_form_data.php", "form bag\u3001form_load", "\u5f9e bag\uff0fload \u79fb\u9664"],
            ["_list.php", "\u52a0\u5217", "\u522a\u5217\uff1b\u5fc5\u8981\u6642\u6539 style.css"],
            ["\u8cc7\u6599\u5eab", "ALTER TABLE \u52a0\u6b04", "\u522a\u6b04\uff08\u6ce8\u610f\u65e2\u6709\u8cc7\u6599\uff09"],
        ],
    )


def section_3(doc):
    heading(doc, "\u4e09\u3001\u6ce8\u610f\u4e8b\u9805", level=1)
    notes = [
        "_config.php \u662f\u8d77\u9ede\uff1a\u8868\u540d\u3001\u5916\u9375\u3001CSRF key \u932f\u4e86\u6703\u9023\u9396\u5931\u6557\u3002",
        "CSRF key \u4e0d\u53ef\u91cd\u8907\uff1a\u6bcf\u55ae\u5143 csrf\u3001list_csrf \u61c9\u552f\u4e00\u3002",
        "\u7d22\u5f15\u9801\uff08\u7b2c\u56db\u7ae0\uff09\u8207\u8868\u55ae\u9801\uff08\u7b2c\u4e94\u7ae0\uff09\u5206\u958b\u7dad\u8b77\uff1b\u5217\u8868\u6539 _list.php\uff0c\u8868\u55ae\u6539 _detail.php\uff0faddin.php\u3002",
        "\u65b0\u589e\uff0f\u4fee\u6539\u5171\u7528 _detail.php\uff1b\u50c5\u8cc7\u6599\u6e96\u5099\u5206\u5728 add.php \u8207 update.php\u3002",
        "\u5171\u7528\u6b04\u4f4d \u2192 _inc.php\uff08\u7b2c\u516d\u7ae0\uff09\uff1b\u81ea\u8a02\u6b04\u4f4d \u2192 \u55ae\u5143 _detail.php\uff0faddin.php\u3002",
        "\u522a\u9664\u5b50\u8868\uff1a\u4e3b\u8868 crud_cfg() cascade _img\uff08\u542b\u6a94\u6848\uff09\u3001_lang\u3001_msg\uff1b\u7279\u6b8a\u5b50\u8868\uff08\u5982 question\uff09\u9700 beforeDelete\u3002",
        "\u8907\u88fd\u8cc7\u6599\u5217\uff1a\u5217\u8868\u8907\u88fd\u6309\u9215 \u2192 add.php?PKey=\uff1b\u5716\u7247\u901a\u5e38\u4e0d\u8907\u88fd\u3002",
        "\u5b50\u6a21\u7d44\u512a\u5148\u7528 manage_child_list_prepare\uff0fmanage_child_form_*\u3002",
        "_list.php \u52a0\u6b04\u5f8c\u8868\u683c\u8981\u5728 style.css \u88dc .tableGrid--*\u3002",
        "\u65b0\u589e Module_PKey \u5f8c\u8981\u66f4\u65b0 _inc.php \u7684 $manage_module_detail_fields\uff08\u7b2c\u516d\u7ae0\uff09\u3002",
        "program.strLink \u5fc5\u9808\u8207 manage/{strLink}/ \u76ee\u9304\u540d\u4e00\u81f4\u3002",
    ]
    for i, n in enumerate(notes, 1):
        add_para(doc, f"{i}. {n}")

    heading(doc, "\u5feb\u901f\u5c0d\u7167\uff1a\u4f9d\u9700\u6c42\u6539\u54ea\u4e9b\u6a94", level=2)
    add_table(
        doc,
        ["\u9700\u6c42", "\u4e3b\u8981\u4fee\u6539"],
        [
            ["\u5168\u65b0\u5167\u5bb9\u55ae\u5143", "\u8907\u88fd\u76ee\u9304 10 \u6a94\uff0bDB\uff0bmodule_p\uff0b_inc.php\uff08\u7b2c\u516d\u7ae0\uff09"],
            ["\u53ea\u958b\u95dc\u9996\u9801\uff0f\u6a19\u7c64", "_inc.php\uff08\u7b2c\u516d\u7ae0\uff09"],
            ["\u8868\u55ae\u52a0\u81ea\u8a02\u6b04", "_detail.php + addin.php + _form_data.php\uff08\u7b2c\u4e94\u7ae0\u5171\u7528\u5340\uff09"],
            ["\u5217\u8868\u591a\u4e00\u6b04", "_list.php + style.css\uff08\u7b2c\u56db\u7ae0\u7d22\u5f15\u5340\uff09"],
            ["\u52a0\u5b50\u6a21\u7d44", "\u8907\u88fd album_d \u6216 question_class + \u7236 _list.php"],
        ],
    )


def section_4(doc):
    heading(doc, "\u56db\u3001\u7d22\u5f15\u9801\u9762\uff08list\uff09\u8aaa\u660e", level=1)
    add_para(
        doc,
        "\u7d22\u5f15\u9801\u662f\u5f8c\u53f0\u55ae\u5143\u7684\u5217\u8868\u5165\u53e3\uff0c\u8207\u65b0\u589e\uff0f\u4fee\u6539\u8868\u55ae\u9801\u5206\u96e2\u3002"
        "\u4fee\u6539\u7d22\u5f15\u9801\u6642\uff0c\u61c9\u907f\u514d\u6539\u5230\u8868\u55ae\u5132\u5b58\u908f\u8f2f\uff08addin.php\uff09\u3002",
    )

    heading(doc, "4.1 \u7d22\u5f15\u9801\u6a94\u6848\u8207\u8072\u660e", level=2)
    add_table(
        doc,
        ["\u6a94\u6848", "\u6027\u8cea", "\u8072\u660e"],
        [
            ["list.php", "\u55ae\u5143\u5165\u53e3\uff08\u63a7\u5236\u5668\uff09", "CSRF\u3001\u67e5\u8a62\u689d\u4ef6\u3001\u5206\u9801\u3001\u522a\u9664\uff0f\u6392\u5e8f\u8655\u7406\u3001\u8f09\u5165 _list.php"],
            ["_list.php", "\u55ae\u5143\u8996\u5716\uff08View\uff09", "\u8868\u683c\u5217\u3001\u6309\u9215\u3001tableGrid class\uff1b\u4e0d\u5beb SQL \u8207 POST \u8655\u7406"],
            ["_upload.php", "\u5217\u8868\u4e0a\u4e0b\u67b6 API", "\u7531 list \u8868\u55ae data-upload-url \u6307\u5411"],
        ],
    )

    heading(doc, "4.2 \u7d22\u5f15\u9801\u5171\u7528\u5340\uff08manage/ \u6839\u76ee\u9304\uff09", level=2)
    add_para(doc, "list.php \u901a\u5e38 require \u4ee5\u4e0b\u5171\u7528\u6a94\uff0c\u65b0\u55ae\u5143\u4e00\u822c\u4e0d\u9700\u6539\u52d5\uff1a", bold=True)
    add_bullets(
        doc,
        [
            "_inc.php\uff0f_module.php \u2014 \u8a9e\u7cfb\u3001\u6b0a\u9650\u3001Module_PKey",
            "_layout_head.php \u2192 _in_code_head.php \u2192 _in_javascript.php \u2014 \u9801\u9996",
            "_layout_body_open.php \u2192 _header.php \u2192 _sidebar.php \u2014 \u9801\u6846\u958b\u59cb",
            "_breadcrumbs.php \u2014 \u9eb5\u5305\u5c51",
            "_search.php \u2014 \u5206\u985e\u7be9\u9078\uff08\u4f9d Layer \u986f\u793a Class1\u20134\uff09",
            "_select.php \u2014 \u5de5\u5177\u5217\uff08\u5168\u9078\u3001\u522a\u9664\u3001\u6392\u5e8f\u3001\u65b0\u589e\u6309\u9215\uff09",
            "_page.php \u2014 \u5206\u9801",
            "_layout_body_close.php \u2192 _in_code_bottom.php \u2014 \u9801\u6846\u7d50\u675f",
        ],
    )

    heading(doc, "4.3 \u7d22\u5f15\u9801\u5c08\u5c6c\u4fee\u6539\u9ede\uff08\u55ae\u5143\u76ee\u9304\u5167\uff09", level=2)
    add_table(
        doc,
        ["\u6a94\u6848", "\u4fee\u6539\u5167\u5bb9"],
        [
            ["list.php", "list_csrf\u3001table_name\u3001FKName\u3001crud_cfg()\u3001\u641c\u5c0b\u6b04\uff08Keywords\uff09\u3001crud_module_where()"],
            ["_list.php", "\u5217\u51fa\u6b04\u4f4d\u3001\u958b\u5408\u5217\u3001\u7de8\u8f2f\uff0f\u8907\u88fd\u9023\u7d50\u3001manage_module_show_detail_field() \u63a7\u5236\u5217\u986f\u793a"],
            ["manage/css/style.css", "\u65b0\u589e\uff0f\u8abf\u6574\u5217\u6642\u8868\u683c\u5bec\u5ea6\uff1a.tableGrid--{單元名}"],
        ],
    )

    heading(doc, "4.4 \u5b50\u6a21\u7d44\u7d22\u5f15\u9801\u5dee\u7570", level=2)
    add_para(
        doc,
        "\u5b50\u6a21\u7d44\uff08album_d\u3001question_class \u7b49\uff09list.php \u6539\u7528 manage_child_list_prepare()\u3001"
        "manage_child_list_render()\uff0c\u5171\u7528 _child_list_shell.php \u5305\u88dd\u3002"
        "\u55ae\u5143\u5c08\u5c6c\u4ecd\u53ea\u6539 _list.php \u8207 list.php \u7684 prepare \u9078\u9805\u3002",
    )


def section_5(doc):
    heading(doc, "\u4e94\u3001\u65b0\u589e\u8207\u4fee\u6539\u5171\u7528\u5340\u8aaa\u660e", level=1)
    add_para(
        doc,
        "\u65b0\u589e\uff08add.php\uff09\u8207\u4fee\u6539\uff08update.php\uff09\u5171\u7528\u540c\u4e00\u5957\u8868\u55ae\u8996\u5716\u8207\u5132\u5b58\u7aef\u9ede\uff0c"
        "\u50c5\u5728\u5165\u53e3\u6a94\u5206\u5de5\u8cc7\u6599\u6e96\u5099\u9636\u6bb5\u3002",
    )

    heading(doc, "5.1 \u4e09\u5c64\u6a94\u6848\u5206\u5de5", level=2)
    add_table(
        doc,
        ["\u5c64\u7d1a", "\u6a94\u6848", "\u8072\u660e"],
        [
            ["\u5165\u53e3\uff08\u5206\u96e2\uff09", "add.php", "\u521d\u59cb\u5316\u7a7a\u767d\u6216\u8907\u88fd\u4f86\u6e90\uff1bUpdate_PKey=0\uff1b\u4e0d\u5f37\u5236\u6aa2\u67e5\u5217\u5b58\u5728"],
            ["\u5165\u53e3\uff08\u5206\u96e2\uff09", "update.php", "\u5fc5\u9808\u6709\u6709\u6548 PKey\uff1bmanage_detail_set_config($cfg, true) \u9650\u5b9a\u7de8\u8f2f\u7bc4\u570d\uff1b\u8f09\u5165\u4e0d\u5230\u5247\u5c0e\u56de"],
            ["\u5171\u7528\u8996\u5716", "_detail.php", "HTML \u8868\u55ae\u3001fieldCheck JS\u3001\u5716\u7247\u69fd\uff1b\u4ee5 $WorkFile \u5224\u65b7 $isAdd"],
            ["\u5171\u7528\u5132\u5b58", "addin.php", "POST \u9a57\u8b49\u3001\u5beb\u5165\u4e3b\u8868\uff0f\u5b50\u8868\uff0f\u4e0a\u50b3\uff1b\u65b0\u589e\u8207\u4fee\u6539\u540c\u4e00\u652f"],
            ["\u5171\u7528\u8cc7\u6599", "_form_data.php", "form bag\u3001form_init\u3001form_load\u3001form_save\u3001\u8907\u88fd\u908f\u8f2f"],
        ],
    )

    heading(doc, "5.2 add.php \u8207 update.php \u5dee\u7570\u5c0d\u7167", level=2)
    add_table(
        doc,
        ["\u9805\u76ee", "add.php", "update.php"],
        [
            ["PKey / Update_PKey", "\u8907\u88fd\u6642\u53ef\u8b80\u4f86\u6e90 PKey\uff0c\u5132\u5b58\u6642\u4ecd\u70ba 0\uff08INSERT\uff09", "\u5fc5\u9808\u6709 URL\uff0fPOST \u4e4b PKey\uff08UPDATE\uff09"],
            ["manage_detail_set_config", "manage_detail_set_config($cfg)", "manage_detail_set_config($cfg, true) \u4e3b\u6a94\u7bc4\u570d\u6aa2\u67e5"],
            ["form_load", "\u53ef\u9078\uff1a\u8907\u88fd\u6587\u5b57\uff1b\u5716\u7247\u901a\u5e38\u6e05\u7a7a\u91cd\u50b3", "\u5fc5\u9808 form_load \u6210\u529f\uff0c\u5426\u5247\u300c\u67e5\u7121\u8981\u4fee\u6539\u8cc7\u6599\u300d"],
            ["Sort", "\u53ef crud_next_sort() \u81ea\u52d5\u6392\u5e8f", "\u8f09\u5165\u539f\u503c"],
            ["\u9eb5\u5305\u5c51", "manage_breadcrumbs_for_form('\u65b0\u589e')", "manage_breadcrumbs_for_form('\u7de8\u8f2f')"],
            ["\u6700\u7d42 require", "_detail.php\uff08\u5171\u7528\uff09", "_detail.php\uff08\u5171\u7528\uff09"],
        ],
    )

    heading(doc, "5.3 _detail.php \u5167\u5224\u65b7\u65b0\u589e\uff0f\u4fee\u6539", level=2)
    add_para(doc, "\u5178\u578b\u5beb\u6cd5\uff1a$isAdd = stripos((string)($WorkFile ?? ''), 'add') !== false;", bold=True)
    add_bullets(
        doc,
        [
            "\u65b0\u589e\uff1a\u53ef\u9690\u85cf\u4e0a\u50b3\u5716\u9810\u89bd\u3001\u986f\u793a\u7a7a\u767d\u69fd\u4f4d",
            "\u4fee\u6539\uff1a\u986f\u793a\u73fe\u6709\u5716\u3001dtUDate\u3001\u522a\u5716\u6309\u9215\uff08_del_img.php\uff09",
            "\u8868\u55ae action \u7d71\u4e00\u6307\u5411 addin.php\uff1b\u532f\u51fa\u6642\u7531 addin \u5224\u65b7 INSERT \u6216 UPDATE",
            "\u6b04\u4f4d\u986f\u793a\u4ecd\u7528 manage_module_show_detail_field()\uff0c\u8207\u7b2c\u516d\u7ae0 _inc.php \u8a2d\u5b9a\u9023\u52d5",
        ],
    )

    heading(doc, "5.4 addin.php \u5171\u7528\u5132\u5b58\u6d41\u7a0b", level=2)
    add_bullets(
        doc,
        [
            "crud_csrf_verify_form($csrfKey) \u2014 CSRF \u8207 _config csrf \u5c0d\u61c9",
            "\u5f9e filter_array \u53d6 formPKey\uff1b<=0 \u70ba\u65b0\u589e\uff1b>0 \u70ba\u4fee\u6539",
            "\u9a57\u8b49\u6b04\u4f4d\uff08\u9806\u5e8f\u3001\u8a9e\u7cfb\u3001\u5206\u985e\u3001\u5716\u7247\u5927\u5c0f\u7b49\uff09",
            "\u5beb\u5165\u4e3b\u8868 \u2192 \u5b50\u8868 lang\uff0fmsg\uff0fimg \u2192 \u95dc\u806f\u6a19\u7c64\uff08\u82e5\u6709\uff09",
            "crud_addin_return_url() \u5c0e\u56de update \u6216 list",
        ],
    )

    heading(doc, "5.5 \u65b0\u589e\uff0f\u4fee\u6539\u5171\u7528\u6a94\u6e05\u55ae", level=2)
    add_table(
        doc,
        ["\u6a94\u6848", "\u662f\u5426\u5171\u7528", "\u8abf\u6574\u6b04\u4f4d\u6642\u4fee\u6539\u9ede"],
        [
            ["_config.php", "\u662f", "\u8868\u540d\u3001csrf\u3001photo_slots\u3001tag \u7b49"],
            ["_form_data.php", "\u662f", "form bag \u9810\u8a2d\u503c\u3001load/save\u3001\u8907\u88fd\u8655\u7406"],
            ["_detail.php", "\u662f", "HTML \u6b04\u4f4d\u3001JS fieldCheck\u3001$isAdd \u5206\u652f"],
            ["addin.php", "\u662f", "\u9a57\u8b49\u8207 DB \u5beb\u5165"],
            ["_upload.php", "\u662f", "\u4e0a\u4e0b\u67b6\u5207\u63db\uff08\u82e5\u6709\uff09"],
            ["_del_img.php", "\u662f", "\u7de8\u8f2f\u9801\u522a\u5716\uff08\u82e5\u6709\u5716\uff09"],
            ["add.php", "\u5426\uff08\u50c5\u65b0\u589e\uff09", "\u8907\u88fd\u4f86\u6e90\u3001\u521d\u59cb Sort"],
            ["update.php", "\u5426\uff08\u50c5\u4fee\u6539\uff09", "PKey \u6aa2\u67e5\u3001form_load \u5931\u6557\u5c0e\u56de"],
            ["list.php / _list.php", "\u5426", "\u7d22\u5f15\u9801\u7368\u7acb\uff0c\u898b\u7b2c\u56db\u7ae0"],
        ],
    )

    heading(doc, "5.6 \u5b50\u6a21\u7d44\u65b0\u589e\uff0f\u4fee\u6539\u5171\u7528", level=2)
    add_table(
        doc,
        ["\u4e00\u822c\u5167\u5bb9\u55ae\u5143", "\u5b50\u6a21\u7d44\uff08manage_child_*\uff09"],
        [
            ["add.php \u81ea\u5beb\u521d\u59cb\u5316", "manage_child_form_add_prepare()"],
            ["update.php \u81ea\u5beb load", "manage_child_form_update_prepare()"],
            ["addin.php \u81ea\u5beb", "manage_child_addin_run() \u6216 _form_data \u5167 validate/save"],
            ["_detail.php \u5171\u7528", "_detail.php \u5171\u7528\uff08\u8207 add/update \u540c\u6a23\uff09"],
        ],
    )

    heading(doc, "5.7 \u4fee\u6539\u6b04\u4f4d\u6642\u5feb\u901f\u5224\u65b7\u6539\u54ea\u88e1", level=2)
    add_table(
        doc,
        ["\u9700\u6c42", "\u7d22\u5f15 list", "\u65b0\u589e/\u4fee\u6539\u5171\u7528", "\u50c5 add", "\u50c5 update"],
        [
            ["\u5217\u8868\u591a\u4e00\u6b04", "_list.php + style.css", "\u2014", "\u2014", "\u2014"],
            ["\u8868\u55ae\u591a\u4e00\u6b04", "\u2014", "_detail.php + addin.php + _form_data.php", "\u2014", "\u2014"],
            ["\u8907\u88fd\u5217\u884c\u908f\u8f2f", "\u2014", "_form_data.php", "add.php \u8b80\u8907\u88fd PKey", "\u2014"],
            ["\u7de8\u8f2f\u524d\u6aa2\u67e5\u7236\u9375", "\u2014", "addin verify_edit_row", "\u2014", "update.php form_load"],
            ["\u522a\u9664\u3001\u6392\u5e8f", "list.php", "\u2014", "\u2014", "\u2014"],
        ],
    )


def section_6(doc):
    heading(doc, "\u516d\u3001manage/_inc.php\uff1a$manage_module_detail_fields \u8a2d\u5b9a\u8aaa\u660e", level=1)
    add_para(
        doc,
        "\u6a94\u6848\u4f4d\u7f6e\uff1amanage/_inc.php\uff08\u7d04\u7b2c 276 \u884c\u8d77\uff09\u3002"
        "\u7528\u65bc\u540c\u4e00\u5957\u8868\u55ae\u7a0b\u5f0f\u4f9d\u300c\u55ae\u5143\u300d\u958b\u95dc\u5171\u7528\u9078\u7528\u6b04\u4f4d\uff0c"
        "\u7121\u9808\u70ba\u6bcf\u500b module_p \u8907\u88fd _detail.php\u3002",
    )

    heading(doc, "6.1 \u914d\u7f6e\u7d50\u69cb", level=2)
    add_table(
        doc,
        ["\u9375", "\u610f\u7fa9"],
        [
            ["'_default'", "\u672a\u5217\u51fa\u7684 Module_PKey \u5957\u7528\u4e4b\u9810\u8a2d\u503c\uff08\u901a\u5e38\u5168 false\uff09"],
            ["\u6578\u5b57\u9375\uff081\u30017\u30018\u2026\uff09", "module_p \u8cc7\u6599\u8868 PKey\uff0c\u5373\u7db2\u5740 manNo\uff0f\u5168\u57df $Module_PKey"],
        ],
    )

    heading(doc, "6.2 \u6b04\u4f4d\u9375\uff08$field\uff09\u5c0d\u7167", level=2)
    add_table(
        doc,
        ["\u9375", "UI\uff0fDB", "\u8aaa\u660e"],
        [
            ["home", "Home \u6b04\uff1b\u5217\u8868\u300c\u9996\u9801\u300d\u6b04", "\u9996\u9801\u5340\u5854\u662f\u5426\u5448\u73fe"],
            ["interview", "Interview1\u3001Interview2\u2026", "\u591a\u8a9e\u7cfb\u7c21\u8ff0 textarea"],
            ["list", "Photo1 \u5217\u8868\u5716", "\u5167\u5bb9\u660e\u7d30\u9801\uff08paper\u3001news\u3001product \u7b49\uff09"],
            ["list_class", "Photo1 \u5217\u8868\u5716", "\u5206\u985e\u6a21\u7d44 class1/2/3\uff1b\u53ef\u8207 list \u5206\u958b\u8a2d"],
            ["tag", "tag_d \u95dc\u806f", "\u9808 _config \u8a2d tag_relation_parent_col \u4e14 addin \u6709\u5132\u5b58\u908f\u8f2f"],
        ],
    )

    heading(doc, "6.3 \u5224\u65b7\u51fd\u5f0f", level=2)
    add_para(doc, "manage_module_show_detail_field(string $field): bool", bold=True)
    add_bullets(
        doc,
        [
            "$field \u53ef\u70ba\uff1ahome\u3001interview\u3001list\u3001list_class\u3001tag",
            "\u4ee5 $Module_PKey\uff08\u7531 _module.php \u5f9e manNo \u8f09\u5165\uff09\u67e5\u8868\uff1b\u7121\u5c0d\u61c9\u9375\u5247\u7528 _default",
            "\u55ae\u4e00\u9375\u672a\u5b9a\u7fa9\u6642\uff0c\u56de\u5239 _default \u8a72\u9375",
        ],
    )

    heading(doc, "6.4 \u65b0\u589e\u55ae\u5143\u6642\u8a2d\u5b9a\u6b65\u9a5f", level=2)
    add_bullets(
        doc,
        [
            "\u5f8c\u53f0\u300c\u55ae\u5143\u8a2d\u5b9a\u300d\u5efa\u7acb module_p\uff0c\u8a18\u4e0b PKey\uff08\u5373 list.php?manNo=\u6b64\u503c\uff09",
            "\u5728 $manage_module_detail_fields \u65b0\u589e\uff1aPKey => ['home'=>bool, 'interview'=>bool, 'list'=>bool, 'list_class'=>bool, 'tag'=>bool]",
            "\u8907\u88fd\u55ae\u5143\u76ee\u9304\u6642\uff0c_detail.php \u82e5\u5df2\u6709 manage_module_show_detail_field()\u3001\u6703\u81ea\u52d5\u751f\u6548",
            "tag \u8a2d true \u6642\uff0c\u78ba\u8a8d _config.php \u6709 tag_relation_parent_col\uff0c\u4e14 addin.php \u6709\u5c0d\u61c9\u5132\u5b58",
        ],
    )

    heading(doc, "6.5 \u7a0b\u5f0f\u5f15\u7528\u7bc4\u4f8b", level=2)
    add_table(
        doc,
        ["\u6a94\u6848", "\u5178\u578b\u5beb\u6cd5"],
        [
            ["_detail.php", "$showHomeField = manage_module_show_detail_field('home');\n<?php if ($showHomeField) { ?> \u2026 <?php } ?>"],
            ["_list.php", "$showHomeColumn = manage_module_show_detail_field('home');\n\u5217\u8868\u9996\u9801\u6b04\u4f9d\u6b64\u986f\u793a"],
            ["addin.php", "if (manage_module_show_detail_field('tag')) { \u2026 \u5132\u5b58 tag_d \u2026 }"],
            ["class1/_detail.php", "manage_module_show_detail_field('list_class') \u2014 \u5206\u985e\u7528 list_class \u800c\u975e list"],
        ],
    )

    heading(doc, "6.6 \u73fe\u884c\u7bc4\u4f8b\u914d\u7f6e\uff08\u7d22\u5f15\uff09", level=2)
    add_para(doc, "\u4ee5\u4e0b\u70ba\u672c\u7ad9\u5df2\u8a2d\u5b9a\u4e4b module_p.PKey\uff08\u55ae\u5143\u540d\u7a31\u4ee5\u5f8c\u53f0\u55ae\u5143\u8a2d\u5b9a\u70ba\u6e96\uff09\uff1a")
    add_table(
        doc,
        ["PKey", "home", "interview", "list", "list_class", "tag"],
        [
            ["7", "true", "true", "true", "false", "true"],
            ["8", "true", "true", "true", "false", "true"],
            ["9", "false", "false", "true", "false", "false"],
            ["10", "true", "true", "true", "false", "true"],
            ["13", "false", "false", "true", "false", "false"],
            ["17", "false", "false", "true", "false", "false"],
            ["18", "false", "false", "true", "false", "false"],
            ["20", "false", "false", "true", "false", "false"],
        ],
    )

    heading(doc, "6.7 \u6ce8\u610f\u4e8b\u9805", level=2)
    add_bullets(
        doc,
        [
            "\u50c5\u63a7\u5236\u5171\u7528\u9078\u7528\u6b04\uff1b\u55ae\u5143\u81ea\u8a02\u6b04\u4f4d\u8acb\u6539\u8a72\u55ae\u5143 _detail.php\uff0faddin.php",
            "\u6b64\u8655 true \u4f46\u8868\u55ae\u672a\u5305 manage_module_show_detail_field \u8005\uff0c\u6b04\u4f4d\u4ecd\u4e0d\u986f\u793a\uff08\u9700\u5169\u908a\u4e00\u81f4\uff09",
            "addin \u5132\u5b58\u908f\u8f2f\u82e5\u672a\u5224\u65b7\u958b\u95dc\uff0c\u53ef\u80fd\u51fa\u73fe\u300c\u7121\u8868\u55ae\u537b\u5beb\u5165\u300d",
            "list \u8207 list_class \u53ef\u5206\u958b\uff1a\u660e\u7d30\u6709\u5716\u3001\u5206\u985e\u5217\u8868\u7121\u5716",
            "\u65b0\u55ae\u5143\u672a\u5217\u5165\u6642\u884c\u70ba\u540c _default\uff08\u591a\u6578\u6b04\u4f4d\u9810\u8a2d\u4e0d\u986f\u793a\uff09",
        ],
    )


def section_7(doc):
    heading(doc, "\u4e03\u3001.env \u74b0\u5883\u8b8a\u6578\u8a2d\u5b9a\u5c0d\u61c9\u8aaa\u660e", level=1)
    add_para(
        doc,
        "\u8b8a\u6578\u7531 include/host.php \u65bc\u7cfb\u7d71\u555f\u52d5\u6642\u8f09\u5165\uff08manage/_inc.php \u7b2c\u4e00\u884c require\uff09\u3002"
        "\u672c\u6a5f WAMP \u53ef\u653e\u65bc\u5c08\u6848\u6839\u76ee\u9304 brick6/.env\uff1b"
        "\u6b63\u5f0f\u6a5f\u5efa\u8b70\u6539\u7528 config/env.path.php \u6307\u5411 private \u76ee\u9304\uff0c\u907f\u514d .env \u66b4\u9732\u65bc httpdocs\u3002",
    )

    heading(doc, "7.1 \u6a94\u6848\u4f4d\u7f6e\u8207\u8f09\u5165\u9806\u5e8f", level=2)
    add_bullets(
        doc,
        [
            "1\uff09\u5c08\u6848\u6839 brick6/.env\uff08\u958b\u767c\u6a5f\u5e38\u7528\uff09",
            "2\uff09config/env.path.php \u2192 \u6307\u5b9a\u975e\u516c\u958b\u76ee\u9304\uff08\u6b63\u5f0f\u6a5f\uff09",
            "3\uff09\u5411\u4e0a\u641c\u5c0b\uff1a.env\u3001config/env.local.php\u3001Plesk private/.env",
            "4\uff09\u5099\u7528\uff1ainclude/.env",
            "\u5df2\u5b58\u5728\u7684\u7cfb\u7d71\u74b0\u5883\u8b8a\u6578\u512a\u5148\uff1b\u6a94\u6848\u5167 # \u70ba\u8a3b\u89e3",
        ],
    )

    heading(doc, "7.2 \u61c9\u7528\u74b0\u5883", level=2)
    add_table(
        doc,
        ["\u8b8a\u6578\u540d", "\u793a\u7bc4\uff0f\u9810\u8a2d", "\u5c0d\u61c9\u8aaa\u660e"],
        [
            ["APP_ENV", "production", "development\uff1flocal \u986f\u793a PHP \u932f\u8aa4\uff1bproduction \u95b1\u85cf\u756b\u9762\u932f\u8aa4\uff08host.php app_configure_error_display\uff09"],
            ["APP_SECRET_KEY", "\u2014", "\u52a0\u5bc6\u7bc7\u7528\u5bc6\u9470\uff08Function.php\uff09\uff1b\u672a\u8a2d\u6642\u6709\u9810\u8a2d\u503c\uff0c\u6b63\u5f0f\u6a5f\u5fc5\u586b"],
        ],
    )

    heading(doc, "7.3 \u7db2\u57df\u3001HTTPS \u8207 Proxy", level=2)
    add_table(
        doc,
        ["\u8b8a\u6578\u540d", "\u793a\u7bc4\uff0f\u9810\u8a2d", "\u5c0d\u61c9\u8aaa\u660e"],
        [
            ["APP_DOMAIN", "example.com", "\u4e3b\u7db2\u57df\uff1b\u7528\u65bc RootURL\u3001Cookie \u5efa\u8b70\u503c\uff08APP_DOMAIN \u5e38\u6578\uff09"],
            ["TRUSTED_DOMAINS", "host1,host2", "\u5141\u8a31\u7684 Host \u767d\u540d\u55ae\uff08\u9017\u865f\u3001\u7a7a\u767d\u5206\u9694\uff09\uff1b\u672a\u8a2d\u6642\u4ee5\u7576\u524d\u8acb\u6c42 Host \u70ba\u6e96"],
            ["NO_HTTPS_DOMAINS", "tsg12", "\u5217\u51fa\u4e0d\u5f37\u5236 HTTPS \u7684\u4e3b\u6a5f\uff08\u672c\u6a5f\u958b\u767c\u7528\uff09\uff1b\u5f71\u97ff current_scheme() \u8207 HSTS"],
            ["FORCE_HTTPS", "1", "1\uff0ftrue \u6642\u5c07 http \u5c0e\u5411 https\uff1b0 \u95dc\u9589\u5f37\u5236"],
            ["TRUSTED_PROXIES", "127.0.0.1", "CDN\uff0f\u4f4d\u5143\u4e4b\u4fe1\u4efb IP \u6216 CIDR\uff1b\u624d\u6703\u63a5\u53d7 X-Forwarded-* \u6a19\u982d"],
        ],
    )

    heading(doc, "7.4 \u8cc7\u6599\u5eab\uff08Conn.php\uff09", level=2)
    add_table(
        doc,
        ["\u8b8a\u6578\u540d", "\u793a\u7bc4\uff0f\u9810\u8a2d", "\u5c0d\u61c9\u8aaa\u660e"],
        [
            ["DB_HOST", "localhost", "MySQL \u4e3b\u6a5f"],
            ["DB_PORT", "3306", "MySQL \u9023\u7dda\u57e0"],
            ["DB_NAME", "\u2014", "\u8cc7\u6599\u5eab\u540d\u7a31\uff08\u5fc5\u586b\uff09"],
            ["DB_USER", "\u2014", "\u767b\u5165\u5e33\u865f\uff08\u5fc5\u586b\uff09"],
            ["DB_PASSWD", "\u2014", "\u767b\u5165\u5bc6\u78bc\uff08\u5fc5\u586b\uff09"],
            ["DB_SOCKET", "\u2014", "Unix socket \u8def\u5f91\uff08\u53ef\u9078\uff0c\u6709\u503c\u6642\u512a\u5148\u65bc host:port\uff09"],
            ["DB_TIMEZONE", "\u2014", "PDO \u9023\u7dda\u5f8c SET time_zone\uff08\u53ef\u9078\uff09"],
            ["DB_SCHEMA_CACHE_TTL", "300", "\u8cc7\u6599\u8868\u7d50\u69cb\u5feb\u53d6\u79d2\u6578\uff08dbclass.php\uff09"],
            ["MYSQL_SSL_CA / CERT / KEY / CAPATH / CIPHER", "\u2014", "MySQL SSL \u9023\u7dda\uff08\u96f2\u7aef\u53ef\u9078\uff09"],
            ["MYSQL_SSL_VERIFY", "1", "1 \u9a57\u8b49\u4f3a\u8b49\uff1b0 \u95dc\u9589"],
        ],
    )

    heading(doc, "7.5 \u5b89\u5168\u8207 Cookie", level=2)
    add_table(
        doc,
        ["\u8b8a\u6578\u540d", "\u793a\u7bc4\uff0f\u9810\u8a2d", "\u5c0d\u61c9\u8aaa\u660e"],
        [
            ["PASSWORD_PEPPER", "\u2014", "\u5bc6\u78bc\u96dc\u8a0a\uff08HMAC \u5f8c\u518d password_hash\uff09\uff1b\u672a\u8a2d\u6642\u767b\u5165\u76f8\u95dc\u51fd\u5f0f\u6703\u62cb\u932f\uff08\u5fc5\u586b\uff09"],
            ["APP_ENABLE_HSTS", "1", "1 \u9001 Strict-Transport-Security\uff08Conn.php\uff09"],
            ["APP_HSTS_MAX_AGE", "31536000", "HSTS \u6709\u6548\u79d2\u6578\uff08\u9810\u8a2d 1 \u5e74\uff09"],
            ["APP_HSTS_INCLUDE_SUBDOMAINS", "1", "HSTS \u662f\u5426\u542b\u5b50\u7db2\u57df"],
            ["APP_HSTS_PRELOAD", "0", "HSTS preload \u6a19\u793a\uff081\uff0f0\uff09"],
        ],
    )

    heading(doc, "7.6 \u7b2c\u4e09\u65b9\u670d\u52d9", level=2)
    add_table(
        doc,
        ["\u8b8a\u6578\u540d", "\u793a\u7bc4\uff0f\u9810\u8a2d", "\u5c0d\u61c9\u8aaa\u660e"],
        [
            ["RECAPTCHA_SITE_KEY", "\u2014", "\u5f8c\u53f0\u767b\u5165\u9801\u8868\u55ae\u986f\u793a\uff08manage/login/index.php\uff09"],
            ["RECAPTCHA_SECRET", "\u2014", "\u4f3a\u670d\u5668\u7aef\u9a57\u8b49 reCAPTCHA\uff1b\u7a7a\u503c\u5247\u4e0d\u555f\u7528"],
            ["MAIL_API_URL", "http://webmail...\u3002php", "\u767c\u4fe1 API \u4f4d\u5740\uff08Function.php \u767c\u4fe1\u51fd\u5f0f\uff09"],
            ["SMS_USERNAME", "\u2014", "\u7c21\u8a0a\u767c\u9001\u5e33\u865f\uff08Function.php\uff0c\u6709\u7528\u5230\u624d\u586b\uff09"],
            ["SMS_PASSWORD", "\u2014", "\u7c21\u8a0a\u767c\u9001\u5bc6\u78bc"],
        ],
    )

    heading(doc, "7.7 \u8207\u5f8c\u53f0\u55ae\u5143\u958b\u767c\u95dc\u806f", level=2)
    add_bullets(
        doc,
        [
            "web_root\u3001web_url \u7531 host.php \u4f9d\u76ee\u9304\u8207 APP_DOMAIN \u8a08\u7b97\uff0c\u4e00\u822c\u4e0d\u5beb\u5728 .env",
            "Upload \u5be6\u9ad4\u76ee\u9304\u70ba\u5c08\u6848\u4e0b Upload/\uff08UPLOAD_BASE\uff09\uff0c\u8207 .env \u7121\u95dc",
            "\u65b0\u589e\u55ae\u5143\u4e0d\u9700\u6539 .env\uff1b\u65b0\u7db2\u57df\u3001\u65b0\u6a5f\u5668\u6216\u65b0\u8cc7\u6599\u5eab\u624d\u9700\u8abf\u6574",
            ".env \u542b\u5bc6\u78bc\u8207\u5bc6\u9470\uff0c\u52ff\u63d0\u4ea4\u7248\u672c\u5eab\uff1b.gitignore \u61c9\u6392\u9664 .env",
        ],
    )

    heading(doc, "7.8 \u672c\u7ad9 .env \u5df2\u4f7f\u7528\u6b04\u4f4d\u7d22\u5f15", level=2)
    add_para(doc, "\u4ee5\u4e0b\u70ba brick6/.env \u5be6\u969b\u5b58\u5728\u7684\u8b8a\u6578\u540d\uff08\u503c\u8acb\u5728\u6a5f\u5668\u4e0a\u76f4\u63a5\u7de8\u8f2f\uff0c\u4e0d\u5beb\u5165\u6587\u4ef6\uff09\uff1a")
    add_table(
        doc,
        ["\u5206\u985e", "\u8b8a\u6578\u540d"],
        [
            ["\u8cc7\u6599\u5eab", "DB_HOST, DB_PORT, DB_USER, DB_PASSWD, DB_NAME"],
            ["\u61c9\u7528", "APP_SECRET_KEY, APP_ENV"],
            ["\u7db2\u57df\u8207 HTTPS", "TRUSTED_DOMAINS, APP_DOMAIN, TRUSTED_PROXIES, FORCE_HTTPS, NO_HTTPS_DOMAINS"],
            ["\u5b89\u5168", "PASSWORD_PEPPER"],
            ["\u767b\u5165\u9a57\u8b49", "RECAPTCHA_SITE_KEY, RECAPTCHA_SECRET"],
            ["\u90f5\u4ef6", "MAIL_API_URL"],
        ],
    )


def section_8(doc):
    heading(doc, "八、前端程式的使用設定", level=1)
    add_para(
        doc,
        "前端頁面使用後台單元資料時，Module_PKey 統一由 include/frontend_modules.php 管理，"
        "程式以 frontend_module_pkey('slug') 取得，不在頁面中直接寫數字。"
        "前端共用函式位於 include/frontend_helpers.php，並已由根目錄 _inc.php 載入。",
    )

    heading(doc, "8.1 相關檔案與用途", level=2)
    add_table(
        doc,
        ["檔案", "用途"],
        [
            ["include/frontend_modules.php", "前端 slug 與 module_p.PKey 的唯一對照表（主要設定檔）"],
            ["include/frontend_helpers.php", "提供 frontend_module_pkey()、列表、明細、分類、麵包屑等共用函式"],
            ["/_inc.php", "載入 frontend_helpers.php，並建立 $Array_MU_Name、$Array_MU_Link 等前台選單資料"],
            ["manage/{單元}/_config.php", "共用後台資料表設定；前端可 array_merge() 後補 view、連結、排序與分頁設定"],
            ["/{單元}.php、/{單元}-detail.php", "前端列表頁與明細頁；依 slug 取得 Module_PKey 並呼叫共用函式"],
        ],
    )

    heading(doc, "8.2 新增前端單元的設定步驟", level=2)
    add_bullets(
        doc,
        [
            "先完成後台 module_p 建立與單元程式設定，確認實際 module_p.PKey。",
            "在 include/frontend_modules.php 增加 'slug' => PKey；slug 使用穩定、可辨識的英文小寫名稱。",
            "前端頁面先 require('_inc.php')，再以 $Module_PKey = frontend_module_pkey('slug'); 取值。",
            "若頁面使用 frontend_fetch_list()、frontend_fetch_detail() 等共用資料函式，呼叫 frontend_module_set_config() 設定資料表、view、連結、排序與分頁。",
            "以 $Array_MU_Name[$Module_PKey] 與 $Array_MU_Link[$Module_PKey] 取得後台單元名稱及前台連結，並確認列表、明細與分類頁均可正常開啟。",
        ],
    )

    heading(doc, "8.3 Module_PKey 取值方式", level=2)
    add_table(
        doc,
        ["使用情境", "建議寫法", "注意事項"],
        [
            ["已登錄的固定單元", "$Module_PKey = frontend_module_pkey('news');", "slug 必須存在於 frontend_modules.php；未定義會拋出 RuntimeException"],
            ["依後台 PageLink 動態反查", "$Module_PKey = frontend_module_pkey_for_link('events.htm');", "僅用於尚未登錄或確實需依連結識別的單元；查無資料回傳 0，必須自行檢查"],
            ["取得名稱與連結", "$Module_Name = $Array_MU_Name[$Module_PKey] ?? '';\n$Module_Link = $Array_MU_Link[$Module_PKey] ?? $page_link;", "需先載入 _inc.php，避免直接重查 module_p/module_lang"],
        ],
    )
    add_para(doc, "禁止寫法：$Module_PKey = 8;（PKey 因資料移轉或重建而變動時，頁面會讀錯單元）", bold=True)

    heading(doc, "8.4 前端共用模組設定範例", level=2)
    add_para(
        doc,
        "$Module_PKey = frontend_module_pkey('news');\n"
        "frontend_module_set_config(array_merge(\n"
        "    require __DIR__ . '/manage/news/_config.php',\n"
        "    [\n"
        "        'view' => 'view_news',\n"
        "        'class_link' => 'news',\n"
        "        'detail_link' => 'news-detail',\n"
        "        'publish_window' => true,\n"
        "        'order_by' => 'OpenDate DESC',\n"
        "        'page_size' => 12,\n"
        "        'class1_filter_min_count' => 2,\n"
        "    ]\n"
        "));",
    )
    add_table(
        doc,
        ["設定鍵", "用途／預設"],
        [
            ["master、fk、img、lang、msg、link", "沿用 manage/{單元}/_config.php 的主表、外鍵及子表設定；master、fk 必須有效"],
            ["view", "前端查詢 view；預設 view_{master}"],
            ["class_link", "分類連結前綴；預設 master"],
            ["detail_link", "明細連結前綴；預設 {master}-detail"],
            ["publish_window", "true 時加入 OpenDate／EndDate 上下架期間條件；預設 false"],
            ["order_by", "列表排序；只接受英數字、底線、逗號及空白，無效時回退 PKey DESC"],
            ["page_size", "每頁筆數；預設 12，最小 1"],
            ["class1_filter_min_count", "分類數達此門檻才顯示分類篩選；預設 2，最小 1"],
        ],
    )

    heading(doc, "8.5 列表頁／明細頁基本順序", level=2)
    add_table(
        doc,
        ["頁型", "基本流程"],
        [
            ["列表頁", "設定 $pageName／$subPageName → require _inc.php → 取得 Module_PKey → set_config → 名稱／連結與麵包屑 → 分類條件 → total／paginate → fetch_list"],
            ["明細頁", "設定 $pageName／$subPageName → require _inc.php → 取得 Module_PKey 與 PKey → set_config → frontend_fetch_detail() → 查無資料時回列表或顯示 404 → 載入圖片／語系／相關資料"],
            ["首頁區塊", "直接以 frontend_module_pkey('slug') 放入查詢參數；仍不得寫死 PKey 數字"],
        ],
    )

    heading(doc, "8.6 現行 frontend_modules.php 對照", level=2)
    add_para(doc, "以下為目前設定；實際值以 include/frontend_modules.php 為準：")
    add_table(
        doc,
        ["slug", "module_p.PKey", "單元"],
        [
            ["knowledge", "7", "專欄主題"],
            ["news", "8", "最新消息"],
            ["product", "10", "工業產品"],
            ["company", "12", "關於我們"],
            ["faq", "13", "相關問題"],
            ["filedown", "14", "檔案下載"],
            ["investor", "15", "投資人專區"],
            ["weblink", "17", "相關網站"],
            ["video", "18", "影音專區"],
            ["question", "19", "問卷調查"],
            ["album", "20", "相簿"],
            ["contact", "21", "聯絡我們"],
        ],
    )

    heading(doc, "8.7 注意事項", level=2)
    add_bullets(
        doc,
        [
            "module_p.PKey 異動時，只更新 frontend_modules.php 的對照值；再全站搜尋 frontend_module_pkey('slug') 驗證使用處。",
            "slug 一旦使用便不要任意改名；若必須改名，需同步修改所有列表、明細、首頁區塊與測試程式。",
            "frontend_module_set_config() 的 master、fk、view 必須是合法識別字；格式無效會拋出 RuntimeException，不存在則會在查詢時失敗。",
            "publish_window 僅適用具有 OpenDate、EndDate 的查詢 view；不具欄位的單元保持 false。",
            "frontend_module_pkey_for_link() 依 _inc.php 產生的 $Array_MU_Link 反查；PageLink 拼字、附檔名或後台設定不一致時會得到 0。",
            "前端輸出名稱、網址及資料內容時，仍應使用 e_attr()、safe_url() 等既有安全函式，不因共用設定而省略輸出編碼。",
        ],
    )


def section_9(doc):
    heading(doc, "九、前端版面與安全標頭（MDN HTTP Observatory）", level=1)
    add_para(
        doc,
        "前台頁面除資料查詢外，須遵守共用版型與安全標頭規範。"
        "掃描工具：[MDN HTTP Observatory](https://developer.mozilla.org/en-US/observatory)（目標 A 級；"
        "常見扣分為 HSTS 未送出、Cookie 缺 SameSite、CSP 過寬）。"
        "詳見 MDN 安全標頭說明："
        "https://developer.mozilla.org/en-US/docs/Web/HTTP/Headers/Strict-Transport-Security",
    )

    heading(doc, "9.1 前台共用版型檔", level=2)
    add_table(
        doc,
        ["檔案", "用途", "修改時機"],
        [
            ["_inc.php", "載入 host.php、Conn.php、frontend_helpers；呼叫 send_security_headers()", "新單元需麵包屑變數、body class 時"],
            ["_header.php", "導覽列、次選單（view_dbclass1）", "新增主選單項目、dropdown 邏輯"],
            ["_footer.php", "頁尾聯絡資訊、社群、地圖", "webset 欄位對應、文案"],
            ["_banner.php", "首頁輪播（view_dbad）／內頁 pgBanner + 麵包屑", "新內頁 $pageName 對應 banner 圖"],
            ["_in_code_head.php", "title、meta、OG、JSON-LD", "SEO 覆寫規則"],
            ["_in_javascript.php / _in_code_bottom.php", "CSS/JS 載入順序", "新增外掛時須評估 CSP"],
            ["_code_lang.php", "前台多語字串", "新區塊文案"],
            ["_page_number.php", "列表分頁 HTML", "分頁樣式"],
        ],
    )

    heading(doc, "9.2 新前端頁面檢查清單", level=2)
    add_bullets(
        doc,
        [
            "頁首設定 $pageName（如 p2）、$subPageName；require('_inc.php') 後再查資料。",
            "使用 frontend_module_pkey('slug')，禁止寫死 Module_PKey 數字。",
            "麵包屑 $bread_name / $break_link 與 JSON-LD $ldjson 成對維護。",
            "圖片路徑優先 webp（Upload 目錄 is_file 判斷）；輸出用 e_attr()、safe_href()。",
            "富文本內容用 frontend_render_html()，勿直接 echo 未淨化 HTML。",
            "表單 hidden 使用 name=\"csrf_token\"（與 mail.php 一致）；錯誤集中 #formErrorArea。",
            "含 reCAPTCHA 的頁面：contact.php、events-detail.php（JoinForm=Yes）；金鑰讀 .env。",
            "inline JS 必須 script_open() / script_close() 或 manage_inline_script()，以帶 CSP nonce。",
            "勿在 <head> 加 meta referrer（已改由 Referrer-Policy 標頭統一）。",
        ],
    )

    heading(doc, "9.3 HTTP 安全標頭（include/sec.php）", level=2)
    add_table(
        doc,
        ["項目", "實作位置", "注意事項"],
        [
            ["CSP", "_inc.php → send_security_headers()", "script 用 nonce + strict-dynamic；style 允許 unsafe-inline（相容 WOW/inline style）"],
            ["HSTS", "Plesk/nginx 伺服器層 或 Conn.php（APP_ENABLE_HSTS=1）擇一", "不可重複；tsg5.com.tw 由 Plesk 送 max-age=15768000"],
            ["Referrer-Policy", ".htaccess + sec.php", "strict-origin-when-cross-origin"],
            ["X-Content-Type-Options", ".htaccess", "nosniff"],
            ["Cookie", ".htaccess edit Set-Cookie", "HttpOnly; Secure; SameSite=Lax（HTTPS 環境）"],
            ["X-Frame-Options", "sec.php", "前台 DENY；含 Google 地圖 iframe 時勿加 COEP require-corp"],
        ],
    )

    heading(doc, "9.4 MDN Observatory 常見扣分與對策", level=2)
    add_table(
        doc,
        ["掃描項目", "常見原因", "對策"],
        [
            ["HSTS（-20）", "與 Plesk/nginx 重複送出兩個 Strict-Transport-Security", "伺服器已送 HSTS 時：.env 設 APP_ENABLE_HSTS=0，且 .htaccess 勿再 set HSTS；部署後 curl -sI 應只剩一個標頭"],
            ["CSP", "style-src unsafe-inline、缺 form-action", "前台已設 form-action 'self'；inline style 為相容性保留，Observatory 可能仍提示"],
            ["Cookies", "缺 SameSite", ".htaccess 已加 SameSite=Lax"],
            ["SRI", "外部 script 無 integrity", "站內 script 可加 SRI（sri_manifest.php）；Google reCAPTCHA 勿加 SRI（腳本會更新）"],
            ["Referrer / XCTO", "—", "已通過則維持現狀"],
        ],
    )

    heading(doc, "9.5 .env 與部署（前台相關）", level=2)
    add_bullets(
        doc,
        [
            "APP_ENV=production、FORCE_HTTPS=1（正式機）；TRUSTED_DOMAINS 列入對外網域。",
            "NO_HTTPS_DOMAINS 僅列本機開發用主機名，勿列入正式網域。",
            "TRUSTED_PROXIES 設 CDN/反向代理 IP，否則 HTTPS 判斷失敗導致 HSTS/CSP upgrade-insecure-requests 異常。",
            "RECAPTCHA_SITE_KEY / RECAPTCHA_SECRET：前台聯絡與活動報名。",
            "Laravel 平行站 brick6-laravel：DocumentRoot 指 public/；LEGACY_ASSET_URL 指舊站靜態資源。",
        ],
    )

    heading(doc, "9.6 禁止事項", level=2)
    add_bullets(
        doc,
        [
            "禁止在 _in_code_head.php 使用 referrer meta no-referrer-when-downgrade。",
            "禁止對 Google reCAPTCHA / gstatic 腳本加 SRI integrity（更新後會被瀏覽器封鎖）。",
            "禁止在 .htaccess 拼錯 includeSubdomains（正確：includeSubDomains）。",
            "禁止表單 CSRF 欄位名 csrf 與 csrf_token 混用（mail 端已雙欄位相容，新頁請統一 csrf_token）。",
            "禁止未經 script nonce 的 inline <script>（違反 CSP，Observatory 與瀏覽器主控台均會警告）。",
        ],
    )


def section_10(doc):
    heading(doc, "十、2026/6/26 後台新增與調整", level=1)
    add_para(
        doc,
        "本章記錄 2026 年 6 月 26 日前後完成之後台功能與程式調整，"
        "含側邊選單資料庫化、CSP 相容改造、語系 SEO 三欄位、新聞刊登日期搜尋等。"
        "部署正式站前請一併確認資料庫欄位與檔案清單。",
    )

    heading(doc, "10.1 側邊選單改為資料庫驅動", level=2)
    add_para(
        doc,
        "manage/_sidebar.php 由示範用靜態 $sidebarMenu 改為與 manage/_subNav.php 相同之 MySQL 邏輯，"
        "依 module_p／module_d 與 $_SESSION['FunctionID'] 產生選單；Admin 帳號可見全部項目。",
    )
    add_table(
        doc,
        ["項目", "說明"],
        [
            ["權限", "FunctionID 逗號字串轉整數陣列；非 Admin 僅顯示有權限之 module_p.PKey"],
            ["群組", "網站管理 PKey [4,5,6]、首頁管理 [1,2]、系統管理 [3]；其餘為「單元管理」（intType=1）"],
            ["子選單", "module_d 有資料時為 DROPDOWN；連結 ../{strLink}/list.php?manNo=&subNo="],
            ["外部連結", "urlLink 非空時以 target=_blank、rel=noopener noreferrer 輸出"],
            ["CSP 相容", "側欄收合改 data-manage-action=\"toggle-sidebar\"，由 manage/js/manage-csp.js 委派處理"],
        ],
    )

    heading(doc, "10.2 後台 CSP 與 inline 腳本整合", level=2)
    add_para(
        doc,
        "後台 script-src 不使用 unsafe-inline，改以 nonce + strict-dynamic；"
        "style 仍允許 unsafe-inline（相容舊表單 inline style）。",
    )
    add_table(
        doc,
        ["檔案", "變更要點"],
        [
            ["include/sec.php", "send_manage_security_headers()、script_open()／script_close()、manage_alert_script()、manage_inline_script()"],
            ["manage/_inc.php", "載入後呼叫 send_manage_security_headers()；閒置登出改 manage_alert_script()"],
            ["manage/.htaccess", "移除靜態 CSP，改由 PHP 動態帶 nonce"],
            ["manage/_in_javascript.php、manage/_in_code_bottom.php", "所有 <script src> 加上 nonce 屬性"],
            ["manage/js/manage-csp.js", "新增：側欄、列表全選、外部連結確認、排序按鈕等 data-manage-action 委派"],
            ["各 manage 頁面", "alert／location 導向改 manage_alert_script()；內嵌腳本改 script_open()／script_close()"],
            ["manage/login/login.php", "登入頁改與內頁相同 appRoot 版型（_header、_sidebar、mainContent）"],
        ],
    )
    add_bullets(
        doc,
        [
            "新頁面須先 require manage/_inc.php，否則無 CSP nonce。",
            "禁止新增 onclick=\"...\" 或 javascript: URL；改用 data-manage-action 或 manage-csp.js 擴充。",
            "內嵌 <script> 必須經 script_open() 或 manage_inline_script() 輸出。",
        ],
    )

    heading(doc, "10.3 語系 SEO 三欄位（共用 partial）", level=2)
    add_para(
        doc,
        "多語系內容單元表單新增／調整 SEO 欄位標籤與結構，並抽出共用 partial 避免各單元重複維護。",
    )
    add_table(
        doc,
        ["欄位", "表單名稱", "資料庫", "說明"],
        [
            ["SEO標題", "Title{n}", "{master}_lang.Title", "顯示於 <title>；maxlength=255"],
            ["SEO內文", "Description{n}", "{master}_lang.Description", "meta description；原「Description」標籤更名"],
            ["SEO關鍵字", "Keyword1_{n}～Keyword5_{n}", "{master}_lang.Keywords", "五個關鍵字輸入框；原「關鍵字」標籤更名"],
        ],
    )
    add_para(doc, "已套用單元（9 個）：news、paper、album、faq、product、video、company、filedown、question", bold=True)
    add_table(
        doc,
        ["檔案", "用途"],
        [
            ["manage/_detail_lang_seo_fields.php", "新建：語系分頁內 SEO 三欄 HTML；父層需提供迴圈變數 $i"],
            ["manage/{單元}/_detail.php", "語系區塊內 require（勿用 require_once，否則英文分頁不顯示）"],
            ["manage/{單元}/_form_data.php", "form_load 後呼叫 manage_lang_apply_seo_from_lang_data($langData)"],
            ["manage/_child_helpers.php", "manage_lang_apply_seo_from_lang_data()：Title → class1_form_vars.SeoTitle"],
            ["include/crud_helpers.php", "crud_load_lang_slots_data／crud_save_lang_slots 讀寫 Title（crud_table_has_column 防呆）"],
        ],
    )
    add_bullets(
        doc,
        [
            "新單元要 SEO 欄位：複製上述三處引用，並確認 _config.php 的 lang 表有 Title 欄位。",
            "partial 內使用 $SeoTitle、$Description、$Keywords 陣列（與既有表單變數一致）。",
        ],
    )

    heading(doc, "10.4 資料庫：*_lang.Title 欄位", level=2)
    add_para(
        doc,
        "SEO 標題儲存於各語系子表 Title 欄位。正式站須先執行 ALTER，後台儲存才會寫入。"
        "程式以 crud_table_has_column() 判斷，未加欄位時不會報錯但 SEO 標題無法保存。",
    )
    add_para(doc, "語系表範例（依各單元 _config.php 的 lang 鍵）：", bold=True)
    add_bullets(
        doc,
        [
            "news_lang、paper_lang、album_lang、faq_lang、product_lang、video_lang、company_lang、filedown_lang、question_lang",
            "dbclass1_lang、dbclass2_lang、dbclass3_lang、dbad_lang、dbweb_lang、tag_lang、module_lang、investor_lang 等",
        ],
    )
    add_para(
        doc,
        "ALTER 範本（請 DBA 依實際表名執行，欄位已存在則略過）：\n"
        "ALTER TABLE `news_lang` ADD COLUMN `Title` VARCHAR(255) NULL COMMENT '標題' AFTER `strName`;",
    )

    heading(doc, "10.5 新聞列表：刊登日期搜尋", level=2)
    add_para(
        doc,
        "manage/news/list.php 篩選區新增「刊登日期（起）」「刊登日期（迄）」，"
        "比對主表 OpenDate 欄位；後端邏輯沿用 include/crud_helpers.php 的 crud_list_apply_opendate_range()，"
        "與 manage/member/list.php（dtDate）相同模式。",
    )
    add_table(
        doc,
        ["項目", "說明"],
        [
            ["表單欄位", "OpenDate、EndDate（type=date）"],
            ["驗證", "起日大於迄日時顯示 listDateRangeError 錯誤訊息"],
            ["清除", "filterWrap 清除按鈕連結 list.php?manNo=&subNo= 重設條件"],
            ["其他單元", "若主表有 OpenDate，可複製 news/list.php 篩選區與 crud_list_apply_opendate_range 呼叫"],
        ],
    )

    heading(doc, "10.6 內容圖版型說明顯示調整", level=2)
    add_para(
        doc,
        "paper、news、company 的 _detail.php 中，四行版型尺寸說明（上圖下文／左圖右文等）"
        "自 for ($n=2...) 內容圖迴圈移出，改在迴圈結束後只顯示一組，避免多語系分頁重複出現四次。",
    )

    heading(doc, "10.7 修改程式清單總表", level=2)
    add_table(
        doc,
        ["類別", "檔案路徑"],
        [
            ["側欄選單", "manage/_sidebar.php"],
            ["CSP／安全", "include/sec.php、manage/_inc.php、manage/.htaccess、manage/js/manage-csp.js"],
            ["CSP 版型", "manage/_in_javascript.php、manage/_in_code_bottom.php、manage/login/login.php"],
            ["SEO 共用", "manage/_detail_lang_seo_fields.php、manage/_child_helpers.php、include/crud_helpers.php"],
            ["SEO 單元 _detail", "manage/news、paper、album、faq、product、video、company、filedown、question/_detail.php"],
            ["SEO 單元 _form_data", "manage/news、paper、album、faq、product、video、company、filedown、question/_form_data.php"],
            ["列表搜尋", "manage/news/list.php"],
            ["圖說明調整", "manage/news/_detail.php、manage/paper/_detail.php、manage/company/_detail.php"],
        ],
    )

    heading(doc, "10.8 部署注意事項", level=2)
    add_bullets(
        doc,
        [
            "上傳上述檔案後以 Ctrl+F5 強制重新整理後台，確認側欄、登入頁與 CSP 主控台無錯誤。",
            "正式站須先執行 *_lang.Title 的 ALTER，再於後台編輯 SEO 標題。",
            "若英文（或其他語系）分頁看不到 SEO 區塊，檢查 _detail.php 是否誤用 require_once 載入 partial。",
            "弱掃相關：根目錄 .htaccess 已阻擋 composer.json、package.json、.env 等直接存取（與後台無關但建議一併部署）。",
        ],
    )


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_doc_title(doc, "\u5f8c\u53f0\u65b0\u589e\uff0f\u8abf\u6574\u55ae\u5143")
    add_doc_subtitle(doc, "\u4fee\u6539\u7a0b\u5f0f\u6e05\u55ae\u8207\u6ce8\u610f\u4e8b\u9805")

    # 章節順序：前言 → 一 → 二 → 三 → 四 → 五 → 六 → 七 → 八 → 九 → 十
    section_preface(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)

    try:
        doc.save(str(OUT))
        print("Saved:", OUT)
    except PermissionError:
        doc.save(str(OUT_ALT))
        print("Saved (alt, close original docx first):", OUT_ALT)


if __name__ == "__main__":
    main()
